import os
import json
import shutil
import subprocess
import tempfile
import threading
import re
from datetime import datetime, timedelta
import tkinter as tk
from tkinter import messagebox, filedialog
import customtkinter as ctk
import pandas as pd

# Library untuk pemrosesan Word
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Library untuk Live Preview PDF
try:
    from docx2pdf import convert as convert_to_pdf_word
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

try:
    import fitz  # PyMuPDF
    from PIL import Image
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

# Library kalender
try:
    from tkcalendar import DateEntry
    HAS_TKCALENDAR = True
except ImportError:
    HAS_TKCALENDAR = False

# Konfigurasi Tema
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

DATA_FILE = "sips_data.json"
HISTORY_FILE = "sips_history.json"
DATABASE_XLSX = "database_dprd_asn.xlsx"

# Daftar jenis surat yang bisa dipilih pada panel Live Preview.
PREVIEW_TEMPLATES = [
    ("Surat Tugas (DPRD)",     "__surat_tugas_dprd__",    "ctx"),
    ("Surat Tugas (ASN)",      "__surat_tugas_asn__",     "ctx"),
    ("Surat Pemberitahuan",    "pemberitahuan_dprd.docx", "ctx"),
    ("SPD DPRD - Halaman Depan",    "SPD_DPRD.docx",           "person_dprd"),
    ("SPD DPRD - Halaman Belakang", "SPD_BELAKANG.docx",       "person_dprd"),
    ("SPD ASN - Halaman Depan",     "SPD_DPRD.docx",           "person_asn"),
    ("SPD ASN - Halaman Belakang",  "SPD_BELAKANG.docx",       "person_asn"),
    ("Daftar Hadir",                "__daftar_hadir__",        "ctx"),
]

# Template Surat Tugas: dipisah antara format "biasa" (<=3 pelaksana) dan
# format "tabel" (>3 pelaksana), masing-masing untuk DPRD dan ASN.
TEMPLATE_ST_DPRD_BIASA = "surat_tugas_dprd_biasa.docx"
TEMPLATE_ST_DPRD_TABEL = "surat_tugas_dprd_tabel.docx"
TEMPLATE_ST_ASN_BIASA = "surat_tugas_asn_biasa.docx"
TEMPLATE_ST_ASN_TABEL = "surat_tugas_asn_tabel.docx"

# Urutan kategori pelaksana DPRD sesuai slot pada surat pemberitahuan
# (pelaksana_tugas_1..4 / jlh_pelaksana_dprd1..4).
KATEGORI_DPRD_ORDER = ["Pimpinan DPRD", "Komisi I", "Komisi II", "Komisi III"]

# Daftar kota/kabupaten di Sulawesi Utara (untuk deteksi transportasi & penentuan tanggal)
SULAWESI_UTARA_CITIES = [
    "Bitung", "Manado", "Tomohon", "Kotamobagu",
    "Minahasa", "Minahasa Utara", "Minahasa Selatan", "Minahasa Tenggara",
    "Bolaang Mongondow", "Bolaang Mongondow Utara", "Bolaang Mongondow Selatan",
    "Bolaang Mongondow Timur", "Kepulauan Sangihe", "Kepulauan Talaud", "Kepulauan Sitaro"
]

# Daftar kota di Jabodetabek (untuk penentuan tujuan awal SPD belakang)
JABODETABEK_CITIES = ["Jakarta", "Bekasi", "Tangerang", "Depok", "Bogor"]

# ===========================================================================
# HELPER: Increment nomor
# ===========================================================================

def increment_nomor(nomor_base, increment=0):
    if increment == 0:
        return nomor_base
    parts = nomor_base.split('/')
    try:
        angka = int(parts[0].strip())
        parts[0] = str(angka + increment)
        return '/'.join(parts)
    except (ValueError, IndexError):
        return nomor_base

def increment_nomor_spd(nomor_base, increment=0):
    if increment == 0:
        return nomor_base
    parts = nomor_base.split('/')
    try:
        angka = int(parts[0].strip())
        parts[0] = str(angka + increment)
        return '/'.join(parts)
    except (ValueError, IndexError):
        return nomor_base

# ===========================================================================
# HELPER: Ekstrak nama kota/kabupaten/provinsi dari string
# ===========================================================================

def extract_city_name(text):
    text = text.strip()
    match = re.search(r'(Kota|Kabupaten|Provinsi)\s+([\w\s]+)', text)
    if match:
        return f"{match.group(1)} {match.group(2).strip()}"
    if "DKI Jakarta" in text:
        return "Kota Jakarta"
    if "Jakarta" in text:
        return "Kota Jakarta"
    return text

def is_in_sulawesi_utara(city_name):
    for c in SULAWESI_UTARA_CITIES:
        if c in city_name:
            return True
    return False

def is_in_jabodetabek(city_name):
    for c in JABODETABEK_CITIES:
        if c in city_name:
            return True
    return False

# ===========================================================================
# HELPER: Generate tanggal dan hari untuk setiap tujuan (tanpa duplikasi hari)
# ===========================================================================

def generate_periods(tanggal_mulai_str, destinations):
    bulan_map = {
        "Januari": 1, "Februari": 2, "Maret": 3, "April": 4,
        "Mei": 5, "Juni": 6, "Juli": 7, "Agustus": 8,
        "September": 9, "Oktober": 10, "November": 11, "Desember": 12
    }
    parts = tanggal_mulai_str.split()
    if len(parts) == 3:
        day = int(parts[0])
        month = bulan_map.get(parts[1], 1)
        year = int(parts[2])
        start_date = datetime(year, month, day)
    else:
        start_date = datetime.now()

    first_city = extract_city_name(destinations[0]) if destinations else ""
    offset = 0 if is_in_sulawesi_utara(first_city) else 1
    base_date = start_date + timedelta(days=offset)

    periods = []
    for idx, dest in enumerate(destinations):
        current_date = base_date + timedelta(days=idx)
        hari_indonesia = {
            "Monday": "Senin", "Tuesday": "Selasa", "Wednesday": "Rabu",
            "Thursday": "Kamis", "Friday": "Jumat", "Saturday": "Sabtu", "Sunday": "Minggu"
        }
        hari_eng = current_date.strftime("%A")
        hari = hari_indonesia.get(hari_eng, hari_eng)
        bulan_indo = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
                      "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        tanggal_str = f"{current_date.day} {bulan_indo[current_date.month]} {current_date.year}"
        periods.append({
            "tujuan": dest,
            "hari": hari,
            "tanggal": tanggal_str
        })
    return periods

# ===========================================================================
# HELPER: Table borders & formatting (tidak berubah)
# ===========================================================================

def set_cell_font(cell, font_name="Arial", font_size=11, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
        pPr = para._p.get_or_add_pPr()
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))
        rPr.append(sz)
        pPr.append(rPr)

def set_table_fixed_widths(table, col_widths_twips):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    total_w = sum(col_widths_twips)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total_w))
    tblW.set(qn('w:type'), 'dxa')

    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_twips):
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(col_widths_twips[i]))
                tcW.set(qn('w:type'), 'dxa')

def set_all_cell_borders(table, border_style="single", border_size=4, border_color="000000"):
    border_xml = {
        'top': (border_style, border_size, border_color),
        'bottom': (border_style, border_size, border_color),
        'left': (border_style, border_size, border_color),
        'right': (border_style, border_size, border_color),
        'insideH': (border_style, border_size, border_color),
        'insideV': (border_style, border_size, border_color),
    }
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)
            tcBorders = OxmlElement('w:tcBorders')
            for side, (style, size, color) in border_xml.items():
                border_el = OxmlElement(f'w:{side}')
                border_el.set(qn('w:val'), style)
                border_el.set(qn('w:sz'), str(size))
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), color)
                tcBorders.append(border_el)
            tcPr.append(tcBorders)

def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def _set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

# ===========================================================================
# HELPER: Surat Tugas dengan tabel bergaris
# ===========================================================================

def _find_table_by_header(doc, header_keywords):
    for tbl in doc.tables:
        if len(tbl.rows) < 2:
            continue
        header_texts = [c.text.strip().lower() for c in tbl.rows[0].cells]
        if all(any(kw.lower() in h for h in header_texts) for kw in header_keywords):
            return tbl
    return None

def _set_cell_text_preserve_style(cell, lines):
    para = cell.paragraphs[0]
    font_name, font_size, font_bold = "Arial", Pt(11), False
    if para.runs:
        r0 = para.runs[0]
        if r0.font.name:
            font_name = r0.font.name
        if r0.font.size:
            font_size = r0.font.size
        font_bold = bool(r0.font.bold)

    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for p in list(cell.paragraphs[1:]):
        p._element.getparent().remove(p._element)

    for idx, line in enumerate(lines):
        if idx > 0:
            run = para.add_run()
            run._element.append(OxmlElement('w:br'))
        run = para.add_run(line)
        run.font.name = font_name
        run.font.size = font_size
        run.bold = font_bold

def _fill_table_rows_from_master(doc, header_keywords, rows_data):
    tbl = _find_table_by_header(doc, header_keywords)
    if tbl is None or not rows_data:
        return

    template_tr = tbl.rows[1]._tr
    parent = template_tr.getparent()
    new_trs = []
    for data_row in rows_data:
        new_tr = copy.deepcopy(template_tr)
        new_trs.append(new_tr)

    ref = template_tr
    for new_tr in new_trs:
        ref.addnext(new_tr)
        ref = new_tr
    parent.remove(template_tr)

    for new_tr, data_row in zip(new_trs, rows_data):
        cells_tc = new_tr.findall(qn('w:tc'))
        from docx.table import _Cell
        for tc, val in zip(cells_tc, data_row):
            cell = _Cell(tc, tbl)
            lines = val.split('\n') if isinstance(val, str) else [str(val)]
            _set_cell_text_preserve_style(cell, lines)

# ===========================================================================
# CLEANUP: Hapus blok kosong pada Surat Tugas Biasa (diperbaiki)
# ===========================================================================

def _remove_empty_blocks(doc, placeholder_prefixes):
    pass

def cleanup_surat_tugas_biasa(doc, template_type):
    paragraphs = list(doc.paragraphs)
    to_remove = set()

    if template_type == 'dprd':
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            stripped = p.text.replace('\t', ' ').strip()
            if re.match(r'^Nama\s*:\s*$', stripped):
                to_remove.add(id(p))
                if i + 1 < len(paragraphs):
                    next_p = paragraphs[i + 1]
                    next_stripped = next_p.text.replace('\t', ' ').strip()
                    if re.match(r'^Jabatan\s*:\s*$', next_stripped):
                        to_remove.add(id(next_p))
                        i += 2
                        continue
            i += 1
    else:
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            stripped = p.text.replace('\t', ' ').strip()
            if re.match(r'^Nama\s*:\s*$', stripped):
                to_remove.add(id(p))
                j = i + 1
                while j < len(paragraphs):
                    ns = paragraphs[j].text.replace('\t', ' ').strip()
                    if re.match(r'^(Pangkat|NIP|Jabatan)\s*:\s*$', ns):
                        to_remove.add(id(paragraphs[j]))
                        j += 1
                    else:
                        break
                i = j
                continue
            i += 1

    for p in paragraphs:
        if id(p) in to_remove:
            p._element.getparent().remove(p._element)

# ===========================================================================
# SURAT TUGAS DPRD & ASN (dengan cleanup)
# ===========================================================================

def buat_surat_tugas_dprd(ctx, selected_dprd, out_path):
    n = len(selected_dprd)
    render_ctx = ctx.copy()

    if n <= 3:
        template_path = TEMPLATE_ST_DPRD_BIASA
        for i in range(1, 4):
            p = selected_dprd[i - 1] if i <= n else {}
            render_ctx[f"pelaksana_tugas_{i}"] = p.get('nama', '')
            render_ctx[f"jabatan_pelaksana_{i}"] = p.get('jabatan', '')
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        cleanup_surat_tugas_biasa(doc, 'dprd')
        doc.save(out_path)
    else:
        template_path = TEMPLATE_ST_DPRD_TABEL
        render_ctx["loop"] = {"index": ""}
        render_ctx["tabel"] = {"nama": "", "jabatan": ""}
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        rows_data = [[str(i + 1), p.get('nama', ''), p.get('jabatan', '')]
                     for i, p in enumerate(selected_dprd)]
        _fill_table_rows_from_master(doc, ["No", "Nama", "Jabatan"], rows_data)
        doc.save(out_path)

def buat_surat_tugas_asn(ctx, selected_asn, out_path):
    n = len(selected_asn)
    render_ctx = ctx.copy()

    if n <= 3:
        template_path = TEMPLATE_ST_ASN_BIASA
        for i in range(1, 4):
            p = selected_asn[i - 1] if i <= n else {}
            render_ctx[f"nama_asn_{i}"] = p.get('nama', '')
            render_ctx[f"pangkat_asn_{i}"] = p.get('pangkat', '')
            render_ctx[f"nip_asn_{i}"] = p.get('nip', '')
            render_ctx[f"jabatan_asn_{i}"] = p.get('jabatan', '')
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        cleanup_surat_tugas_biasa(doc, 'asn')
        doc.save(out_path)
    else:
        template_path = TEMPLATE_ST_ASN_TABEL
        render_ctx["loop"] = {"index": ""}
        render_ctx["tabel"] = {"nama_asn": "", "jabatan_asn": ""}
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        rows_data = []
        for i, p in enumerate(selected_asn):
            nama_col = f"{p.get('nama', '')}\nNIP. {p.get('nip', '-')}"
            jabatan_col = f"{p.get('jabatan', '-')}\n{p.get('pangkat', '-')}"
            rows_data.append([str(i + 1), nama_col, jabatan_col])
        _fill_table_rows_from_master(doc, ["No", "Nama", "Jabatan"], rows_data)
        doc.save(out_path)

# ===========================================================================
# SURAT PEMBERITAHUAN: deteksi otomatis kategori & jumlah pelaksana DPRD
# ===========================================================================

def _label_kategori_dprd(cat, jabatan_list):
    if cat == "Pimpinan DPRD":
        return "Pimpinan DPRD"

    has_pimpinan = any(
        ("ketua" in j.lower() or "sekretaris" in j.lower()) for j in jabatan_list
    )
    has_anggota = any(
        ("anggota" in j.lower() and "ketua" not in j.lower()) for j in jabatan_list
    )

    if has_pimpinan and has_anggota:
        label = f"Pimpinan dan Anggota {cat}"
    elif has_pimpinan:
        label = f"Pimpinan {cat}"
    elif has_anggota:
        label = f"Anggota {cat}"
    else:
        label = cat
    if "DPRD" not in label:
        label += " DPRD"
    return label

def compute_pelaksana_dprd_summary(selected_dprd):
    by_cat = {}
    for p in selected_dprd:
        cat = str(p.get('kategori', '')).strip()
        by_cat.setdefault(cat, []).append(p.get('jabatan', ''))

    summary = []
    for cat in KATEGORI_DPRD_ORDER:
        jabatan_list = by_cat.get(cat)
        if not jabatan_list:
            continue
        label = _label_kategori_dprd(cat, jabatan_list)
        summary.append((label, len(jabatan_list)))

    for cat, jabatan_list in by_cat.items():
        if cat not in KATEGORI_DPRD_ORDER:
            label = _label_kategori_dprd(cat, jabatan_list)
            summary.append((label, len(jabatan_list)))
    return summary

def _remove_empty_pelaksana_lines(doc):
    for p in list(doc.paragraphs):
        txt = p.text
        normalized = re.sub(r'\s+', ' ', txt).strip()

        if "Pendamping ASN" in txt:
            m = re.search(r':\s*(\d+)\s+Orang', normalized)
            if m and int(m.group(1)) == 0:
                p._element.getparent().remove(p._element)
                continue

        if "Kota Bitung" in txt and "Orang" in txt:
            m = re.search(r':\s*(\d+)\s+Orang', normalized)
            if not m:
                p._element.getparent().remove(p._element)
                continue

def apply_pelaksana_dprd_summary_to_ctx(ctx, selected_dprd, max_slots=4):
    summary = compute_pelaksana_dprd_summary(selected_dprd)
    for i in range(1, max_slots + 1):
        if i <= len(summary):
            label, jumlah = summary[i - 1]
            ctx[f"pelaksana_tugas_{i}"] = label
            ctx[f"jlh_pelaksana_dprd{i}"] = jumlah
        else:
            ctx[f"pelaksana_tugas_{i}"] = ""
            ctx[f"jlh_pelaksana_dprd{i}"] = ""
    return ctx

def buat_surat_pemberitahuan_multi(template_path, ctx, selected_dprd, selected_asn, destinations, base_number, out_path, label_asn="Pendamping ASN"):
    periods = generate_periods(ctx.get("tanggal_mulai", ""), destinations)

    base_ctx = ctx.copy()
    apply_pelaksana_dprd_summary_to_ctx(base_ctx, selected_dprd)
    base_ctx["pelaksana_tugas_asn_info"] = label_asn
    base_ctx["jlh_pelaksana_asn"] = len(selected_asn)

    master_doc = None
    first = True

    for idx, period in enumerate(periods):
        nomor_surat = increment_nomor(base_number, idx)
        page_ctx = base_ctx.copy()
        page_ctx["nomor_surat_info"] = nomor_surat
        page_ctx["tujuan_surat_info"] = period["tujuan"]
        page_ctx["hari_info"] = period["hari"]
        page_ctx["tanggal_bertugas_info"] = period["tanggal"]

        tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(page_ctx)
        doc_tpl.save(tmp_docx)

        doc = Document(tmp_docx)
        _remove_empty_pelaksana_lines(doc)
        doc.save(tmp_docx)

        if first:
            master_doc = Document(tmp_docx)
            first = False
        else:
            master_doc.add_page_break()
            subdoc = Document(tmp_docx)
            for element in subdoc.element.body:
                if not element.tag.endswith('sectPr'):
                    master_doc.element.body.append(element)

        os.unlink(tmp_docx)

    if master_doc:
        if master_doc.paragraphs:
            last_p = master_doc.paragraphs[-1]
            if not last_p.text.strip():
                last_p._element.getparent().remove(last_p._element)
        master_doc.save(out_path)

# ===========================================================================
# SPPD: Pisah antara DPRD dan ASN, nomor berbeda
# ===========================================================================

def _build_person_sppd_context(ctx, person, nomor_spd_str, destinations, transport):
    p_ctx = ctx.copy()
    p_ctx["pelaksana_dprd_sppd"] = person.get('nama', '-')
    p_ctx["jabatan_pelaksana_sppd"] = person.get('jabatan', '-')
    p_ctx["nomor_surat_sppd"] = nomor_spd_str
    p_ctx["jenis_perjalanan_sppd"] = ctx.get("jenis_perjalanan", "")
    p_ctx["transportasi_sppd"] = transport

    city_names = [extract_city_name(d) for d in destinations]
    p_ctx["tujuan_bertugas_sppd"] = ", ".join(city_names)

    if any(is_in_jabodetabek(c) for c in city_names):
        tujuan_awal = "Kota Jakarta"
    else:
        tujuan_awal = city_names[0] if city_names else "-"
    p_ctx["tujuan_awal_sppd_belakang"] = tujuan_awal

    p_ctx["tanggal_mulai_sppd"] = ctx.get("tanggal_mulai", "")
    p_ctx["tanggal_akhir_sppd"] = ctx.get("tanggal_akhir", "")
    p_ctx["tanggal_surat_sppd"] = ctx.get("tanggal_surat", "")
    p_ctx["materi_tugas_sppd"] = ctx.get("materi_tugas", "")
    p_ctx["tanggal_mulai_sppd_belakang"] = ctx.get("tanggal_mulai", "")
    p_ctx["tanggal_akhir_sppd_belakang"] = ctx.get("tanggal_akhir", "")
    return p_ctx

def buat_sppd_dprd(spd_depan_template, spd_belakang_template, ctx, sel_dprd, destinations,
                   out_depan, out_belakang):
    import tempfile as tmpmod
    tmpdir = tmpmod.mkdtemp()
    depan_files = []
    belakang_files = []
    nomor_dprd = ctx.get('nomor_spd_dprd', ctx.get('nomor_spd', ''))
    transport = ctx.get("transportasi_otomatis", "Mobil")

    for idx, person in enumerate(sel_dprd):
        p_ctx = _build_person_sppd_context(ctx, person, nomor_dprd, destinations, transport)

        if os.path.exists(spd_depan_template):
            doc_d = DocxTemplate(spd_depan_template)
            doc_d.render(p_ctx)
            tmp = os.path.join(tmpdir, f"dprd_depan_{idx}.docx")
            doc_d.save(tmp)
            depan_files.append(tmp)

        if os.path.exists(spd_belakang_template):
            doc_b = DocxTemplate(spd_belakang_template)
            doc_b.render(p_ctx)
            tmp = os.path.join(tmpdir, f"dprd_belakang_{idx}.docx")
            doc_b.save(tmp)
            belakang_files.append(tmp)

    if depan_files:
        _combine_word_pages(depan_files, out_depan)
    if belakang_files:
        _combine_word_pages(belakang_files, out_belakang)

    shutil.rmtree(tmpdir, ignore_errors=True)

def buat_sppd_asn(spd_depan_template, spd_belakang_template, ctx, sel_asn, destinations,
                  out_depan, out_belakang):
    import tempfile as tmpmod
    tmpdir = tmpmod.mkdtemp()
    depan_files = []
    belakang_files = []
    nomor_asn_base = ctx.get('nomor_spd_asn', ctx.get('nomor_spd', ''))
    transport = ctx.get("transportasi_otomatis", "Mobil")

    for idx, person in enumerate(sel_asn):
        nomor_asn = increment_nomor_spd(nomor_asn_base, idx)
        p_ctx = _build_person_sppd_context(ctx, person, nomor_asn, destinations, transport)

        if os.path.exists(spd_depan_template):
            doc_d = DocxTemplate(spd_depan_template)
            doc_d.render(p_ctx)
            tmp = os.path.join(tmpdir, f"asn_depan_{idx}.docx")
            doc_d.save(tmp)
            depan_files.append(tmp)

        if os.path.exists(spd_belakang_template):
            doc_b = DocxTemplate(spd_belakang_template)
            doc_b.render(p_ctx)
            tmp = os.path.join(tmpdir, f"asn_belakang_{idx}.docx")
            doc_b.save(tmp)
            belakang_files.append(tmp)

    if depan_files:
        _combine_word_pages(depan_files, out_depan)
    if belakang_files:
        _combine_word_pages(belakang_files, out_belakang)

    shutil.rmtree(tmpdir, ignore_errors=True)

def _combine_word_pages(files_list, out_path):
    if not files_list:
        return
    try:
        master = Document(files_list[0])
        for f in files_list[1:]:
            master.add_page_break()
            subdoc = Document(f)
            for element in subdoc.element.body:
                if not element.tag.endswith('sectPr'):
                    master.element.body.append(element)
        if master.paragraphs:
            last_p = master.paragraphs[-1]
            if not last_p.text.strip():
                last_p._element.getparent().remove(last_p._element)
        master.save(out_path)
    except Exception as e:
        print(f"Gagal menggabungkan: {e}")

# ===========================================================================
# APLIKASI UTAMA
# ===========================================================================

class SIPSApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("SIPS - Aplikasi Pembuat Surat Perjalanan Dinas DPRD Bitung")
        self.geometry("1900x900")
        self.minsize(1500, 800)

        self.db_dprd = []
        self.db_asn = []
        self.dprd_vars = {}
        self.asn_vars = {}
        self.pelaksana_vars = {}
        self.pendamping_vars = {}
        self.history_data = {}

        self.mode = "dprd"
        self.current_view = "perjalanan_dinas"
        self.active_categories = {}

        self.database_tujuan = [
            "Kota Manado", "Kota Bitung", "Kota Tomohon", "Kota Kotamobagu",
            "Kabupaten Minahasa", "Kabupaten Minahasa Utara", "Kabupaten Minahasa Selatan",
            "Kabupaten Minahasa Tenggara", "Kabupaten Bolaang Mongondow",
            "Kabupaten Bolaang Mongondow Utara", "Kabupaten Bolaang Mongondow Selatan",
            "Kabupaten Bolaang Mongondow Timur", "Kabupaten Kepulauan Sangihe",
            "Kabupaten Kepulauan Talaud", "Kabupaten Kepulauan Sitaro",
            "DKI Jakarta", "Kota Surabaya", "Kota Bandung", "Kota Medan",
            "Kota Semarang", "Kota Makassar", "Kota Palembang", "Kota Tangerang",
            "Kota Tangerang Selatan", "Kota Bekasi", "Kota Depok", "Kota Yogyakarta",
            "Kota Surakarta (Solo)", "Kota Balikpapan", "Kota Samarinda",
            "Kota Banjarmasin", "Kota Pontianak", "Kota Denpasar", "Kota Mataram",
            "Kota Kupang", "Kota Ambon", "Kota Jayapura", "Kota Sorong",
            "Kota Palu", "Kota Kendari", "Kota Gorontalo", "Kota Palangkaraya",
            "Kota Tarakan", "Kota Banda Aceh", "Kota Padang", "Kota Pekanbaru",
            "Kota Jambi", "Kota Bengkulu", "Kota Bandar Lampung", "Kota Pangkalpinang",
            "Kota Tanjungpinang"
        ]

        self.tujuan_terpilih = []

        self.preview_dir = tempfile.mkdtemp(prefix="sips_preview_")
        self._preview_after_id = None
        self._preview_busy = False
        self._preview_pending = False
        self._preview_lock = threading.Lock()
        self._preview_ctk_image = None

        self.load_database()
        self.load_history()
        self.setup_ui()
        self.calculate_duration()

        self.protocol("WM_DELETE_WINDOW", self.on_close)
        self.after(600, lambda: self.schedule_preview_refresh(immediate=True))

    def on_close(self):
        try:
            shutil.rmtree(self.preview_dir, ignore_errors=True)
        except Exception:
            pass
        self.destroy()

    def terbilang(self, n):
        satuan = ["", "Satu", "Dua", "Tiga", "Empat", "Lima", "Enam", "Tujuh", "Delapan", "Sembilan", "Sepuluh", "Sebelas"]
        if n < 12: return satuan[n]
        elif n < 20: return self.terbilang(n - 10) + " Belas"
        elif n < 100: return self.terbilang(n // 10) + " Puluh " + (satuan[n % 10] if n % 10 != 0 else "")
        return str(n)

    def format_indonesian_date(self, date_obj):
        if not date_obj: return ""
        months = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
                  "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        return f"{date_obj.day} {months[date_obj.month]} {date_obj.year}"

    def normalize_keys(self, data_list):
        normalized = []
        for d in data_list:
            if isinstance(d, dict):
                normalized.append({str(k).lower().strip(): v for k, v in d.items()})
        return normalized

    # ------------------------------------------------------------------
    # DATABASE
    # ------------------------------------------------------------------
    def _read_dprd_asn_from_excel_file(self, path):
        xls = pd.ExcelFile(path)
        sheet_dprd = next((s for s in xls.sheet_names if "dprd" in s.lower() or "anggota" in s.lower()), None)
        sheet_asn = next((s for s in xls.sheet_names if "asn" in s.lower() or "pendamping" in s.lower()), None)

        raw_dprd, raw_asn = [], []
        if sheet_dprd:
            df = pd.read_excel(xls, sheet_name=sheet_dprd)
            df.columns = [str(c).strip() for c in df.columns]
            for _, row in df.iterrows():
                raw_dprd.append({
                    "nama": str(row.get("Nama", row.get("NAMA", ""))).strip(),
                    "jabatan": str(row.get("Jabatan", row.get("JABATAN", ""))).strip(),
                    "kategori": str(row.get("Kategori", row.get("KATEGORI", "Custom"))).strip()
                })
        if sheet_asn:
            df = pd.read_excel(xls, sheet_name=sheet_asn)
            df.columns = [str(c).strip() for c in df.columns]
            for _, row in df.iterrows():
                raw_asn.append({
                    "nama": str(row.get("NAMA", row.get("Nama", ""))).strip(),
                    "nip": str(row.get("NIP", "-")).strip(),
                    "pangkat": str(row.get("PANGKAT/GOLONGAN", row.get("Pangkat", "-"))).strip(),
                    "jabatan": str(row.get("JABATAN", row.get("Jabatan", "-"))).strip()
                })
        return raw_dprd, raw_asn

    def load_database(self):
        if os.path.exists(DATABASE_XLSX):
            try:
                raw_dprd, raw_asn = self._read_dprd_asn_from_excel_file(DATABASE_XLSX)
                if raw_dprd or raw_asn:
                    if raw_dprd: self.db_dprd = self.normalize_keys(raw_dprd)
                    if raw_asn: self.db_asn = self.normalize_keys(raw_asn)
                    self.save_database()
                    return
            except Exception as e:
                print(f"Gagal membaca {DATABASE_XLSX}: {e}")

        if os.path.exists(DATA_FILE):
            try:
                with open(DATA_FILE, "r") as f:
                    data = json.load(f)
                    self.db_dprd = self.normalize_keys(data.get("dprd", []))
                    self.db_asn = data.get("asn", []).copy()
                    return
            except:
                pass

        self.db_dprd = [{"nama": "VIVY JEANET GANAP, S.E.", "jabatan": "KETUA", "kategori": "Pimpinan DPRD"}]
        self.db_asn = [{"nama": "Drs. ALBERT M. SARESE, M.Si.", "nip": "19681011 199010 1 002",
                        "pangkat": "PEMBINA UTAMA MUDA IV/c", "jabatan": "Sekretaris DPRD"}]
        self.save_database()

    def save_database(self):
        try:
            with open(DATA_FILE, "w") as f:
                json.dump({"dprd": self.db_dprd, "asn": self.db_asn}, f, indent=4)
        except:
            pass

    def import_excel_database(self):
        file_path = filedialog.askopenfilename(
            title="Pilih File Database Excel/CSV",
            filetypes=[("Excel & CSV files", "*.xlsx *.xls *.csv")]
        )
        if not file_path: return
        try:
            ext = os.path.splitext(file_path)[1].lower()
            new_dprd, new_asn = [], []
            if ext in (".xlsx", ".xls"):
                new_dprd, new_asn = self._read_dprd_asn_from_excel_file(file_path)
                if not new_dprd and not new_asn:
                    messagebox.showwarning("Format Tidak Dikenali", "Tidak ditemukan sheet yang sesuai.")
                    return
            else:
                df = pd.read_csv(file_path)
                df.columns = [str(c).strip() for c in df.columns]
                cols_lower = [c.lower() for c in df.columns]
                if "nip" in cols_lower:
                    for _, row in df.iterrows():
                        new_asn.append({
                            "nama": str(row.get("NAMA", row.get("Nama", ""))).strip(),
                            "nip": str(row.get("NIP", "-")).strip(),
                            "pangkat": str(row.get("PANGKAT/GOLONGAN", row.get("Pangkat", "-"))).strip(),
                            "jabatan": str(row.get("JABATAN", row.get("Jabatan", "-"))).strip()
                        })
                else:
                    for _, row in df.iterrows():
                        new_dprd.append({
                            "nama": str(row.get("Nama", row.get("NAMA", ""))).strip(),
                            "jabatan": str(row.get("Jabatan", row.get("JABATAN", ""))).strip(),
                            "kategori": str(row.get("Kategori", row.get("KATEGORI", "Custom"))).strip()
                        })

            if new_dprd: self.db_dprd = self.normalize_keys(new_dprd)
            if new_asn: self.db_asn = self.normalize_keys(new_asn)
            self.save_database()
            self.dprd_vars = {}
            self.asn_vars = {}
            self.refresh_personnel_list()
            self.refresh_signer_dropdowns()
            self.schedule_preview_refresh(immediate=True)
            messagebox.showinfo("Import Berhasil",
                f"Database berhasil diperbarui.\nAnggota DPRD: {len(self.db_dprd)} orang\nPendamping ASN: {len(self.db_asn)} orang")
        except Exception as e:
            messagebox.showerror("Import Gagal", f"Terjadi kesalahan:\n{e}")

    def load_history(self):
        if os.path.exists(HISTORY_FILE):
            try:
                with open(HISTORY_FILE, "r") as f:
                    self.history_data = json.load(f)
            except:
                self.history_data = {}
        else:
            self.history_data = {}

    def save_history(self):
        try:
            with open(HISTORY_FILE, "w") as f:
                json.dump(self.history_data, f, indent=4)
        except Exception as e:
            print("Gagal menyimpan riwayat:", e)

    # ------------------------------------------------------------------
    # UI SETUP
    # ------------------------------------------------------------------
    def setup_ui(self):
        self.grid_columnconfigure(0, weight=0, minsize=240)
        self.grid_columnconfigure(1, weight=1, minsize=420)
        self.grid_columnconfigure(2, weight=1, minsize=380)
        self.grid_columnconfigure(3, weight=2, minsize=480)
        self.grid_rowconfigure(0, weight=1)

        self.sidebar_frame = ctk.CTkFrame(self, width=240, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(9, weight=1)

        self.logo_label = ctk.CTkLabel(self.sidebar_frame, text="SIPS DPRD BITUNG",
                                        font=("Arial", 18, "bold"), text_color="#1E3A8A")
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 5))

        self.btn_import_db = ctk.CTkButton(self.sidebar_frame, text="📥 Import Database",
                                            command=self.import_excel_database)
        self.btn_import_db.grid(row=1, column=0, padx=20, pady=10, sticky="ew")

        self.mode_selector = ctk.CTkSegmentedButton(self.sidebar_frame, values=["DPRD", "Setwan"],
                                                     command=self.change_mode)
        self.mode_selector.grid(row=2, column=0, padx=20, pady=10, sticky="ew")
        self.mode_selector.set("DPRD")

        lbl_history = ctk.CTkLabel(self.sidebar_frame, text="Riwayat Edit Surat:", font=("Arial", 11, "bold"))
        lbl_history.grid(row=3, column=0, padx=20, pady=(15, 0), sticky="w")

        history_keys = list(self.history_data.keys())
        self.combo_history = ctk.CTkComboBox(self.sidebar_frame,
                                              values=history_keys if history_keys else ["Tidak ada riwayat"])
        self.combo_history.grid(row=4, column=0, padx=20, pady=(5, 5), sticky="ew")

        self.btn_load_history = ctk.CTkButton(self.sidebar_frame, text="Load Data Surat",
                                               fg_color="#F59E0B", hover_color="#D97706",
                                               command=self.load_selected_history)
        self.btn_load_history.grid(row=5, column=0, padx=20, pady=(0, 15), sticky="ew")

        self.btn_generate_main = ctk.CTkButton(self.sidebar_frame, text="⚡ CETAK SURAT & SPD",
                                                command=self.generate_documents_action,
                                                fg_color="#10B981", hover_color="#059669",
                                                font=("Arial", 14, "bold"))
        self.btn_generate_main.grid(row=6, column=0, padx=20, pady=15, sticky="ew")

        # Kategori Surat Undangan Section
        lbl_undangan_title = ctk.CTkLabel(self.sidebar_frame, text="Kategori Surat Undangan:",
                                          font=("Arial", 11, "bold"))
        lbl_undangan_title.grid(row=7, column=0, padx=20, pady=(10, 5), sticky="w")

        self.btn_undangan_paripurna = ctk.CTkButton(self.sidebar_frame, text="📨 Undangan Paripurna",
                                                     command=self.show_undangan_paripurna,
                                                     fg_color="#6366F1", hover_color="#4F46E5")
        self.btn_undangan_paripurna.grid(row=8, column=0, padx=20, pady=5, sticky="ew")

        self.btn_undangan_biasa = ctk.CTkButton(self.sidebar_frame, text="📋 Undangan Biasa",
                                                command=self.show_undangan_biasa,
                                                fg_color="#8B5CF6", hover_color="#7C3AED")
        self.btn_undangan_biasa.grid(row=9, column=0, padx=20, pady=5, sticky="ew")

        self.btn_back_to_perjalanan = ctk.CTkButton(self.sidebar_frame, text="← Kembali ke Perjalanan Dinas",
                                                     command=self.show_perjalanan_dinas,
                                                     fg_color="#64748B", hover_color="#475569")
        self.btn_back_to_perjalanan.grid(row=10, column=0, padx=20, pady=10, sticky="ew")
        self.btn_back_to_perjalanan.grid_remove()

        self.lbl_credit = ctk.CTkLabel(self.sidebar_frame, text="v7.0 © DPRD Kota Bitung",
                                        font=("Arial", 9), text_color="gray")
        self.lbl_credit.grid(row=11, column=0, padx=20, pady=15, sticky="s")

        # PANEL TENGAH
        self.middle_frame = ctk.CTkScrollableFrame(self, label_text="Data Perjalanan Dinas")
        self.middle_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        self.inputs = {}
        self.mode_specific_widgets = {}
        self.undangan_inputs = {}

        # Nomor Surat
        for var_name, label in [
            ("nomor_surat", "Nomor Surat Tugas DPRD"),
            ("nomor_surat_asn", "Nomor Surat Tugas Setwan"),
        ]:
            lbl = ctk.CTkLabel(self.middle_frame, text=label, anchor="w", font=("Arial", 12, "bold"))
            lbl.pack(fill="x", padx=10, pady=(8, 2))
            ent = ctk.CTkEntry(self.middle_frame, placeholder_text=f"Masukkan {label.lower()}...")
            ent.pack(fill="x", padx=10, pady=(0, 6))
            ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
            self.inputs[var_name] = ent

        # Pemberitahuan DPRD (mode specific)
        self.mode_specific_widgets["lbl_pemberitahuan_dprd"] = ctk.CTkLabel(self.middle_frame, text="Nomor Surat Pemberitahuan DPRD", anchor="w", font=("Arial", 12, "bold"))
        self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(8, 2))
        self.inputs["nomor_pemberitahuan_dprd"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Masukkan nomor surat pemberitahuan dprd...")
        self.inputs["nomor_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(0, 6))
        self.inputs["nomor_pemberitahuan_dprd"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        self.mode_specific_widgets["ent_pemberitahuan_dprd"] = self.inputs["nomor_pemberitahuan_dprd"]

        # Pemberitahuan Setwan
        for var_name, label in [
            ("nomor_pemberitahuan_asn", "Nomor Surat Pemberitahuan Setwan"),
        ]:
            lbl = ctk.CTkLabel(self.middle_frame, text=label, anchor="w", font=("Arial", 12, "bold"))
            lbl.pack(fill="x", padx=10, pady=(8, 2))
            ent = ctk.CTkEntry(self.middle_frame, placeholder_text=f"Masukkan {label.lower()}...")
            ent.pack(fill="x", padx=10, pady=(0, 6))
            ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
            self.inputs[var_name] = ent

        # SPD Section
        lbl_spd_title = ctk.CTkLabel(self.middle_frame, text="Nomor SPD", anchor="w", font=("Arial", 12, "bold"))
        lbl_spd_title.pack(fill="x", padx=10, pady=(10, 2))

        spd_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        spd_frame.pack(fill="x", padx=10, pady=(0, 6))
        spd_frame.grid_columnconfigure(0, weight=1)
        spd_frame.grid_columnconfigure(1, weight=1)

        self.mode_specific_widgets["lbl_spd_dprd"] = ctk.CTkLabel(spd_frame, text="Nomor SPD DPRD :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        self.mode_specific_widgets["lbl_spd_dprd"].grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.inputs["nomor_spd_dprd"] = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 10/SPD/X/2026/")
        self.inputs["nomor_spd_dprd"].grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.inputs["nomor_spd_dprd"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        self.mode_specific_widgets["ent_spd_dprd"] = self.inputs["nomor_spd_dprd"]

        self.mode_specific_widgets["lbl_spd_setwan"] = ctk.CTkLabel(spd_frame, text="Nomor SPD Setwan (ASN) :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        self.mode_specific_widgets["lbl_spd_setwan"].grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.inputs["nomor_spd_asn"] = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 20/SPD/X/2026/")
        self.inputs["nomor_spd_asn"].grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.inputs["nomor_spd_asn"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        self.lbl_spd_pelaksana = ctk.CTkLabel(spd_frame, text="Nomor SPD Pelaksana ASN :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        self.ent_spd_pelaksana = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 10/SPD-PL/X/2026/")
        self.ent_spd_pelaksana.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        self.lbl_spd_pendamping = ctk.CTkLabel(spd_frame, text="Nomor SPD Pendamping ASN :", anchor="w", font=("Arial", 11), text_color="#059669")
        self.ent_spd_pendamping = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 20/SPD-PD/X/2026/")
        self.ent_spd_pendamping.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_spd_info = ctk.CTkLabel(self.middle_frame, text="ℹ️  SPD DPRD: semua anggota pakai nomor sama  |  SPD ASN: nomor otomatis berurutan", anchor="w", font=("Arial", 10), text_color="gray")
        lbl_spd_info.pack(fill="x", padx=10, pady=(0, 6))
        self.mode_specific_widgets["lbl_spd_info"] = lbl_spd_info

        # Dasar Surat Tugas
        lbl_dasar_title = ctk.CTkLabel(self.middle_frame, text="Dasar Surat Tugas", anchor="w", font=("Arial", 12, "bold"))
        lbl_dasar_title.pack(fill="x", padx=10, pady=(10, 2))

        dasar_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        dasar_frame.pack(fill="x", padx=10, pady=(0, 6))
        dasar_frame.grid_columnconfigure(0, weight=1)
        dasar_frame.grid_columnconfigure(1, weight=1)

        lbl_dasar_dprd = ctk.CTkLabel(dasar_frame, text="Dasar Surat Tugas DPRD :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        lbl_dasar_dprd.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.txt_dasar_dprd = ctk.CTkTextbox(dasar_frame, height=60, wrap="word")
        self.txt_dasar_dprd.grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.txt_dasar_dprd.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_dasar_asn = ctk.CTkLabel(dasar_frame, text="Dasar Surat Tugas ASN :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        lbl_dasar_asn.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.txt_dasar_asn = ctk.CTkTextbox(dasar_frame, height=60, wrap="word")
        self.txt_dasar_asn.grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.txt_dasar_asn.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        # Materi / Agenda Kegiatan
        lbl_materi_title = ctk.CTkLabel(self.middle_frame, text="Materi / Agenda Kegiatan", anchor="w", font=("Arial", 12, "bold"))

        materi_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        materi_frame.pack(fill="x", padx=10, pady=(0, 6))
        materi_frame.grid_columnconfigure(0, weight=1)
        materi_frame.grid_columnconfigure(1, weight=1)

        lbl_mt_st = ctk.CTkLabel(materi_frame, text="Surat Tugas & SPPD :", anchor="w",
                                   font=("Arial", 11), text_color="#1E3A8A")
        lbl_mt_st.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.txt_materi_st = ctk.CTkTextbox(materi_frame, height=70, wrap="word")
        self.txt_materi_st.grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.txt_materi_st.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_mt_pb = ctk.CTkLabel(materi_frame, text="Surat Pemberitahuan :", anchor="w",
                                   font=("Arial", 11), text_color="#1E3A8A")
        lbl_mt_pb.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.txt_materi_pb = ctk.CTkTextbox(materi_frame, height=70, wrap="word")
        self.txt_materi_pb.grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.txt_materi_pb.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_jp = ctk.CTkLabel(self.middle_frame, text="Jenis Perjalanan", anchor="w", font=("Arial", 12, "bold"))
        lbl_jp.pack(fill="x", padx=10, pady=(8, 2))
        self.combo_jenis = ctk.CTkComboBox(
            self.middle_frame,
            values=["Kunjungan Kerja", "Kunjungan Konsultasi", "Bimbingan Teknis"],
            command=lambda choice: self.schedule_preview_refresh(immediate=True)
        )
        self.combo_jenis.pack(fill="x", padx=10, pady=(0, 6))

        lbl_tgl_surat = ctk.CTkLabel(self.middle_frame, text="Tanggal Surat", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_surat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_surat = DateEntry(self.middle_frame, width=15, background='#2563EB',
                                       foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_surat.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_surat.bind("<<DateEntrySelected>>", lambda e: self.schedule_preview_refresh())
        else:
            self.dp_surat = ctk.CTkEntry(self.middle_frame)
            self.dp_surat.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_surat.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_tgl_mulai = ctk.CTkLabel(self.middle_frame, text="Tanggal Mulai Tugas", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_mulai.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_mulai = DateEntry(self.middle_frame, width=15, background='#2563EB',
                                       foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_mulai.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_mulai.bind("<<DateEntrySelected>>", lambda e: self.calculate_duration())
        else:
            self.dp_mulai = ctk.CTkEntry(self.middle_frame)
            self.dp_mulai.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_mulai.bind("<KeyRelease>", lambda e: self.calculate_duration())

        lbl_tgl_akhir = ctk.CTkLabel(self.middle_frame, text="Tanggal Akhir Tugas", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_akhir.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_akhir = DateEntry(self.middle_frame, width=15, background='#2563EB',
                                       foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_akhir.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_akhir.bind("<<DateEntrySelected>>", lambda e: self.calculate_duration())
        else:
            self.dp_akhir = ctk.CTkEntry(self.middle_frame)
            self.dp_akhir.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_akhir.bind("<KeyRelease>", lambda e: self.calculate_duration())

        lbl_lama_hari = ctk.CTkLabel(self.middle_frame, text="Lama Hari Perjalanan (Otomatis)",
                                      anchor="w", font=("Arial", 12, "bold"))
        lbl_lama_hari.pack(fill="x", padx=10, pady=(8, 2))
        self.ent_lama_hari = ctk.CTkEntry(self.middle_frame, fg_color="#F3F4F6")
        self.ent_lama_hari.pack(fill="x", padx=10, pady=(0, 6))

        lbl_tujuan = ctk.CTkLabel(self.middle_frame, text="Kota Tujuan Bertugas (Multi Lokasi)",
                                   anchor="w", font=("Arial", 12, "bold"))
        lbl_tujuan.pack(fill="x", padx=10, pady=(8, 2))

        tujuan_input_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        tujuan_input_frame.pack(fill="x", padx=10, pady=(0, 2))
        tujuan_input_frame.grid_columnconfigure(0, weight=1)

        self.ent_tujuan = ctk.CTkEntry(tujuan_input_frame, placeholder_text="Ketik nama kota lalu klik Tambah...")
        self.ent_tujuan.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.ent_tujuan.bind("<KeyRelease>", self.on_tujuan_key_release)
        self.btn_tambah_tujuan = ctk.CTkButton(tujuan_input_frame, text="+ Tambah",
                                                width=80, command=self.tambah_tujuan)
        self.btn_tambah_tujuan.grid(row=0, column=1, sticky="e")

        self.suggestion_frame = ctk.CTkScrollableFrame(self.middle_frame, height=110, fg_color="#F3F4F6")

        lbl_tujuan_terpilih = ctk.CTkLabel(self.middle_frame, text="Lokasi yang dipilih:",
                                             anchor="w", font=("Arial", 11), text_color="gray")
        lbl_tujuan_terpilih.pack(fill="x", padx=10, pady=(4, 1))
        self.tujuan_list_frame = ctk.CTkScrollableFrame(self.middle_frame, height=80, fg_color="#F0F4FF")
        self.tujuan_list_frame.pack(fill="x", padx=10, pady=(0, 6))

        lbl_tujuan_hint = ctk.CTkLabel(self.middle_frame,
            text="ℹ️  Klik ✕ pada lokasi untuk menghapus. Urutan sesuai tampilan.",
            anchor="w", font=("Arial", 10), text_color="gray")
        lbl_tujuan_hint.pack(fill="x", padx=10, pady=(0, 6))

        self.lbl_sign_dprd = ctk.CTkLabel(self.middle_frame, text="Penandatangan DPRD:", font=("Arial", 12, "bold"))
        self.lbl_sign_dprd.pack(fill="x", padx=10, pady=(15, 2))
        self.combo_ttd_dprd = ctk.CTkComboBox(self.middle_frame, values=["-"], height=32,
                                               command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_ttd_dprd.pack(fill="x", padx=10, pady=(0, 8))

        self.lbl_sign_asn = ctk.CTkLabel(self.middle_frame, text="Penandatangan ASN / SPPD:", font=("Arial", 12, "bold"))
        self.lbl_sign_asn.pack(fill="x", padx=10, pady=(10, 2))
        self.combo_ttd_asn = ctk.CTkComboBox(self.middle_frame, values=["-"], height=32,
                                              command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_ttd_asn.pack(fill="x", padx=10, pady=(0, 8))

        self.refresh_signer_dropdowns()

        self.inputs["nomor_surat"].insert(0, "170/DPRD/X/2026")
        self.inputs["nomor_surat_asn"].insert(0, "170/SEK-DPRD/X/2026")
        self.inputs["nomor_pemberitahuan_dprd"].insert(0, "180/DPRD/X/2026")
        self.inputs["nomor_pemberitahuan_asn"].insert(0, "181/SEK-DPRD/X/2026")
        self.inputs["nomor_spd_dprd"].insert(0, "090/SPD/")
        self.inputs["nomor_spd_asn"].insert(0, "091/SPD/")
        self.txt_materi_st.insert("1.0", "Studi Banding terkait Pembahasan Peraturan Daerah")
        self.txt_materi_pb.insert("1.0", "Pimpinan dan Anggota DPRD Kota Bitung akan melakukan Studi Banding terkait Pembahasan Peraturan Daerah")
        self.txt_dasar_dprd.insert("1.0", "Keputusan Pimpinan DPRD Kota Bitung")
        self.txt_dasar_asn.insert("1.0", "Surat Perintah Sekretaris DPRD Kota Bitung")

        self.tujuan_terpilih = ["Kota Manado"]
        self.refresh_tujuan_list_ui()

        # PANEL KANAN
        self.right_frame = ctk.CTkFrame(self)
        self.right_frame.grid(row=0, column=2, padx=10, pady=10, sticky="nsew")
        self.right_frame.grid_columnconfigure(0, weight=1)
        self.right_frame.grid_rowconfigure(4, weight=1)

        lbl_cat_title = ctk.CTkLabel(self.right_frame, text="1. Filter Kategori Calon Pelaksana",
                                      font=("Arial", 13, "bold"), anchor="w")
        lbl_cat_title.grid(row=0, column=0, padx=15, pady=(10, 5), sticky="w")

        self.category_check_frame = ctk.CTkFrame(self.right_frame, fg_color="#F3F4F6")
        self.category_check_frame.grid(row=1, column=0, padx=15, pady=5, sticky="ew")

        self.setup_category_checkboxes()

        lbl_person_title = ctk.CTkLabel(self.right_frame, text="2. Checklist Personel Pelaksana",
                                         font=("Arial", 13, "bold"), anchor="w")
        lbl_person_title.grid(row=2, column=0, padx=15, pady=(15, 2), sticky="w")

        self.btn_action_frame = ctk.CTkFrame(self.right_frame, fg_color="transparent")
        self.btn_action_frame.grid(row=3, column=0, padx=15, pady=2, sticky="ew")
        self.btn_sel_all = ctk.CTkButton(self.btn_action_frame, text="Centang Semua Tampil",
                                          command=self.select_all_visible, height=24)
        self.btn_sel_all.pack(side="left", padx=2)
        self.btn_clear_all = ctk.CTkButton(self.btn_action_frame, text="Bersihkan", fg_color="gray",
                                            command=self.clear_all_visible, height=24)
        self.btn_clear_all.pack(side="left", padx=2)

        self.scroll_personnel = ctk.CTkScrollableFrame(self.right_frame)
        self.scroll_personnel.grid(row=4, column=0, padx=15, pady=(5, 15), sticky="nsew")
        self.refresh_personnel_list()

        # PANEL PREVIEW
        self.preview_frame = ctk.CTkFrame(self)
        self.preview_frame.grid(row=0, column=3, padx=10, pady=10, sticky="nsew")
        self.preview_frame.grid_columnconfigure(0, weight=1)
        self.preview_frame.grid_rowconfigure(3, weight=1)

        lbl_preview_title = ctk.CTkLabel(self.preview_frame, text="👁 Live Preview",
                                          font=("Arial", 13, "bold"), anchor="w")
        lbl_preview_title.grid(row=0, column=0, padx=15, pady=(10, 5), sticky="w")

        self.preview_toolbar = ctk.CTkFrame(self.preview_frame, fg_color="transparent")
        self.preview_toolbar.grid(row=1, column=0, padx=15, pady=(0, 5), sticky="ew")
        self.preview_toolbar.grid_columnconfigure(0, weight=1)

        preview_choices = [item[0] for item in PREVIEW_TEMPLATES]
        self.combo_preview_jenis = ctk.CTkComboBox(
            self.preview_toolbar, values=preview_choices,
            command=lambda choice: self.schedule_preview_refresh(immediate=True)
        )
        self.combo_preview_jenis.set(preview_choices[0])
        self.combo_preview_jenis.grid(row=0, column=0, sticky="ew")

        self.btn_refresh_preview = ctk.CTkButton(
            self.preview_toolbar, text="🔄 Refresh", width=90,
            command=lambda: self.schedule_preview_refresh(immediate=True)
        )
        self.btn_refresh_preview.grid(row=0, column=1, padx=(8, 0))

        self.preview_status_lbl = ctk.CTkLabel(self.preview_frame, text="Menyiapkan preview...",
                                                font=("Arial", 11), text_color="gray",
                                                wraplength=420, justify="left")
        self.preview_status_lbl.grid(row=2, column=0, padx=15, pady=(0, 5), sticky="w")

        self.preview_canvas_frame = ctk.CTkScrollableFrame(self.preview_frame, fg_color="#E5E7EB")
        self.preview_canvas_frame.grid(row=3, column=0, padx=15, pady=(0, 15), sticky="nsew")
        self.preview_image_label = ctk.CTkLabel(self.preview_canvas_frame, text="")
        self.preview_image_label.pack(expand=True, pady=10)

    # ------------------------------------------------------------------
    # METODE UNTUK MODE DPRD / SETWAN
    # ------------------------------------------------------------------
    def change_mode(self, mode):
        self.mode = "setwan" if mode == "Setwan" else "dprd"
        if self.current_view != "perjalanan_dinas":
            self.setup_category_checkboxes()
            self.refresh_personnel_list()
            return
        if self.mode == "dprd":
            # tampilkan dengan pack di posisi yang benar (sebelum label ASN)
            self.lbl_sign_dprd.pack(fill="x", padx=10, pady=(15, 2), before=self.lbl_sign_asn)
            self.combo_ttd_dprd.pack(fill="x", padx=10, pady=(0, 8), before=self.lbl_sign_asn)
            self.inputs["nomor_surat"].configure(placeholder_text="Masukkan nomor surat tugas dprd...")
            # Show DPRD-specific fields
            self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(8, 2))
            self.mode_specific_widgets["ent_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(0, 6))
            self.mode_specific_widgets["lbl_spd_dprd"].grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
            self.mode_specific_widgets["ent_spd_dprd"].grid(row=1, column=0, padx=(0, 4), sticky="ew")
            self.mode_specific_widgets["lbl_spd_setwan"].grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
            self.mode_specific_widgets["lbl_spd_info"].configure(text="ℹ️  SPD DPRD: semua anggota pakai nomor sama  |  SPD ASN: nomor otomatis berurutan")
            # Hide Setwan-specific SPD fields
            self.lbl_spd_pelaksana.grid_forget()
            self.ent_spd_pelaksana.grid_forget()
            self.lbl_spd_pendamping.grid_forget()
            self.ent_spd_pendamping.grid_forget()
            # Update Jenis Perjalanan dropdown for DPRD mode
            self.combo_jenis.configure(values=["Kunjungan Kerja", "Kunjungan Konsultasi", "Bimbingan Teknis"])
        else:
            self.lbl_sign_dprd.pack_forget()
            self.combo_ttd_dprd.pack_forget()
            self.inputs["nomor_surat"].configure(placeholder_text="Masukkan nomor surat tugas setwan...")
            # Hide DPRD-specific fields
            self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack_forget()
            self.mode_specific_widgets["ent_pemberitahuan_dprd"].pack_forget()
            self.mode_specific_widgets["lbl_spd_dprd"].grid_forget()
            self.mode_specific_widgets["ent_spd_dprd"].grid_forget()
            self.mode_specific_widgets["lbl_spd_setwan"].grid_forget()
            self.mode_specific_widgets["lbl_spd_info"].configure(text="ℹ️  SPD Pelaksana & Pendamping: nomor otomatis berurutan per kategori")
            # Show Setwan-specific SPD fields
            self.lbl_spd_pelaksana.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
            self.ent_spd_pelaksana.grid(row=1, column=0, padx=(0, 4), sticky="ew")
            self.lbl_spd_pendamping.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
            self.ent_spd_pendamping.grid(row=1, column=1, padx=(4, 0), sticky="ew")
            # Update Jenis Perjalanan dropdown for Setwan mode
            self.combo_jenis.configure(values=["Studi Komparasi", "Kunjungan Konsultasi", "Bimbingan Teknis"])
        self.setup_category_checkboxes()
        self.category_check_frame.grid(row=1, column=0, padx=15, pady=5, sticky="ew")
        self.refresh_personnel_list()
        self.refresh_signer_dropdowns()
        self.schedule_preview_refresh(immediate=True)

    def setup_category_checkboxes(self):
        for widget in self.category_check_frame.winfo_children():
            widget.destroy()
        self.cat_chk_widgets = {}
        if self.mode == "dprd":
            categories = ["Pimpinan DPRD", "Komisi I", "Komisi II", "Komisi III", "Custom", "Pendamping ASN"]
            default_active = {"Pimpinan DPRD": True, "Komisi I": True, "Komisi II": False,
                              "Komisi III": False, "Custom": False, "Pendamping ASN": True}
        else:
            categories = ["Pelaksana ASN", "Pendamping ASN"]
            default_active = {"Pelaksana ASN": True, "Pendamping ASN": True}
        for idx, cat in enumerate(categories):
            val = tk.BooleanVar(value=default_active.get(cat, False))
            chk = ctk.CTkCheckBox(self.category_check_frame, text=cat, variable=val,
                                   font=("Arial", 11, "bold"), command=self.on_category_changed)
            chk.grid(row=idx // 3, column=idx % 3, padx=10, pady=10, sticky="w")
            self.cat_chk_widgets[cat] = val
        self.active_categories = {cat: var.get() for cat, var in self.cat_chk_widgets.items()}

    def on_category_changed(self):
        for cat, var in self.cat_chk_widgets.items():
            self.active_categories[cat] = var.get()
        self.refresh_personnel_list()

    def refresh_personnel_list(self):
        for widget in self.scroll_personnel.winfo_children():
            widget.destroy()
        self.rendered_dprd_widgets, self.rendered_asn_widgets = {}, {}
        self.rendered_pelaksana_widgets, self.rendered_pendamping_widgets = {}, {}

        if self.mode == "dprd":
            for p in self.db_dprd:
                cat = p.get("kategori", "Custom").strip()
                group_cat = cat if cat in ["Pimpinan DPRD", "Komisi I", "Komisi II", "Komisi III"] else "Custom"
                if self.active_categories.get(group_cat, False):
                    nama, jab = p.get('nama', ''), p.get('jabatan', '')
                    if nama not in self.dprd_vars:
                        self.dprd_vars[nama] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(
                        self.scroll_personnel,
                        text=f"{nama} ({jab}) [{cat}]",
                        variable=self.dprd_vars[nama],
                        command=lambda: self.schedule_preview_refresh(immediate=True)
                    )
                    chk.pack(fill="x", padx=10, pady=4, anchor="w")
                    self.rendered_dprd_widgets[nama] = p

            if self.active_categories.get("Pendamping ASN", False):
                if self.rendered_dprd_widgets:
                    ctk.CTkFrame(self.scroll_personnel, height=2, fg_color="gray").pack(fill="x", padx=10, pady=10)
                for p in self.db_asn:
                    nama, nip, jab = p.get('nama', ''), p.get('nip', '-'), p.get('jabatan', '-')
                    if nama not in self.asn_vars:
                        self.asn_vars[nama] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(
                        self.scroll_personnel,
                        text=f"{nama}\nNIP: {nip} | Jabatan: {jab}",
                        variable=self.asn_vars[nama],
                        command=lambda: self.schedule_preview_refresh(immediate=True)
                    )
                    chk.pack(fill="x", padx=10, pady=5, anchor="w")
                    self.rendered_asn_widgets[nama] = p
        else:
            if self.active_categories.get("Pelaksana ASN", False):
                lbl_pelaksana = ctk.CTkLabel(self.scroll_personnel, text="Pelaksana ASN:",
                                             font=("Arial", 12, "bold"), text_color="#1E3A8A")
                lbl_pelaksana.pack(fill="x", padx=10, pady=(5, 2), anchor="w")
                for p in self.db_asn:
                    nama, nip, jab = p.get('nama', ''), p.get('nip', '-'), p.get('jabatan', '-')
                    key = f"pelaksana_{nama}"
                    if key not in self.pelaksana_vars:
                        self.pelaksana_vars[key] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(
                        self.scroll_personnel,
                        text=f"{nama}\nNIP: {nip} | Jabatan: {jab}",
                        variable=self.pelaksana_vars[key],
                        command=lambda: self.schedule_preview_refresh(immediate=True)
                    )
                    chk.pack(fill="x", padx=10, pady=5, anchor="w")
                    self.rendered_pelaksana_widgets[nama] = p

            if self.active_categories.get("Pendamping ASN", False):
                if self.rendered_pelaksana_widgets:
                    ctk.CTkFrame(self.scroll_personnel, height=2, fg_color="gray").pack(fill="x", padx=10, pady=10)
                lbl_pendamping = ctk.CTkLabel(self.scroll_personnel, text="Pendamping ASN:",
                                              font=("Arial", 12, "bold"), text_color="#059669")
                lbl_pendamping.pack(fill="x", padx=10, pady=(5, 2), anchor="w")
                for p in self.db_asn:
                    nama, nip, jab = p.get('nama', ''), p.get('nip', '-'), p.get('jabatan', '-')
                    key = f"pendamping_{nama}"
                    if key not in self.pendamping_vars:
                        self.pendamping_vars[key] = tk.BooleanVar(value=False)
                    chk = ctk.CTkCheckBox(
                        self.scroll_personnel,
                        text=f"{nama}\nNIP: {nip} | Jabatan: {jab}",
                        variable=self.pendamping_vars[key],
                        command=lambda: self.schedule_preview_refresh(immediate=True)
                    )
                    chk.pack(fill="x", padx=10, pady=5, anchor="w")
                    self.rendered_pendamping_widgets[nama] = p

    def select_all_visible(self):
        for n in self.rendered_dprd_widgets: self.dprd_vars[n].set(True)
        for n in self.rendered_asn_widgets: self.asn_vars[n].set(True)
        for n in self.rendered_pelaksana_widgets:
            key = f"pelaksana_{n}"
            if key in self.pelaksana_vars: self.pelaksana_vars[key].set(True)
        for n in self.rendered_pendamping_widgets:
            key = f"pendamping_{n}"
            if key in self.pendamping_vars: self.pendamping_vars[key].set(True)
        self.schedule_preview_refresh(immediate=True)

    def clear_all_visible(self):
        for n in self.rendered_dprd_widgets: self.dprd_vars[n].set(False)
        for n in self.rendered_asn_widgets: self.asn_vars[n].set(False)
        for n in self.rendered_pelaksana_widgets:
            key = f"pelaksana_{n}"
            if key in self.pelaksana_vars: self.pelaksana_vars[key].set(False)
        for n in self.rendered_pendamping_widgets:
            key = f"pendamping_{n}"
            if key in self.pendamping_vars: self.pendamping_vars[key].set(False)
        self.schedule_preview_refresh(immediate=True)

    # ------------------------------------------------------------------
    # MULTI KOTA TUJUAN
    # ------------------------------------------------------------------
    def tambah_tujuan(self):
        val = self.ent_tujuan.get().strip()
        if not val: return
        if val not in self.tujuan_terpilih:
            self.tujuan_terpilih.append(val)
            self.refresh_tujuan_list_ui()
        self.ent_tujuan.delete(0, tk.END)
        self.hide_tujuan_suggestions()
        self.schedule_preview_refresh(immediate=True)

    def hapus_tujuan(self, kota):
        if kota in self.tujuan_terpilih:
            self.tujuan_terpilih.remove(kota)
            self.refresh_tujuan_list_ui()
            self.schedule_preview_refresh(immediate=True)

    def refresh_tujuan_list_ui(self):
        for w in self.tujuan_list_frame.winfo_children():
            w.destroy()
        if not self.tujuan_terpilih:
            lbl = ctk.CTkLabel(self.tujuan_list_frame, text="(Belum ada tujuan)",
                                text_color="gray", font=("Arial", 11))
            lbl.pack(padx=8, pady=4, anchor="w")
        for kota in self.tujuan_terpilih:
            row_frame = ctk.CTkFrame(self.tujuan_list_frame, fg_color="transparent")
            row_frame.pack(fill="x", padx=4, pady=2)
            lbl = ctk.CTkLabel(row_frame, text=f"📍 {kota}", anchor="w",
                                font=("Arial", 11), text_color="#1E3A8A")
            lbl.pack(side="left", padx=(4, 8))
            btn_del = ctk.CTkButton(row_frame, text="✕", width=28, height=22,
                                     fg_color="#EF4444", hover_color="#DC2626",
                                     font=("Arial", 10, "bold"),
                                     command=lambda k=kota: self.hapus_tujuan(k))
            btn_del.pack(side="right", padx=2)

    def on_tujuan_key_release(self, event):
        val = self.ent_tujuan.get().strip().lower()
        if len(val) >= 2:
            matches = [item for item in self.database_tujuan if val in item.lower()]
            if matches:
                self.show_tujuan_suggestions(matches)
            else:
                self.hide_tujuan_suggestions()
        else:
            self.hide_tujuan_suggestions()
        if event.keysym == "Return":
            self.tambah_tujuan()

    def show_tujuan_suggestions(self, matches):
        for widget in self.suggestion_frame.winfo_children():
            widget.destroy()
        for match in matches[:6]:
            btn = ctk.CTkButton(self.suggestion_frame, text=match, anchor="w",
                                 fg_color="transparent", text_color="black", hover_color="#E5E7EB",
                                 command=lambda m=match: self.select_tujuan_suggestion(m))
            btn.pack(fill="x", padx=5, pady=1)
        self.suggestion_frame.pack(fill="x", padx=10, pady=2, before=self.tujuan_list_frame)

    def hide_tujuan_suggestions(self):
        self.suggestion_frame.pack_forget()

    def select_tujuan_suggestion(self, val):
        self.ent_tujuan.delete(0, tk.END)
        self.ent_tujuan.insert(0, val)
        self.hide_tujuan_suggestions()
        self.tambah_tujuan()

    def refresh_signer_dropdowns(self):
        ttd_dprd_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}"
                           for p in self.db_dprd if str(p.get('kategori', '')).lower() == 'pimpinan dprd']
        self.combo_ttd_dprd.configure(values=ttd_dprd_values if ttd_dprd_values else ["-"])
        self.combo_ttd_dprd.set(ttd_dprd_values[0] if ttd_dprd_values else "-")

        ttd_asn_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}"
                          for p in self.db_asn if 'sekretaris dprd' in str(p.get('jabatan', '')).lower()]
        if not ttd_asn_values:
            ttd_asn_values = [f"{p.get('jabatan', '')} - {p.get('nama', '')}" for p in self.db_asn]
        self.combo_ttd_asn.configure(values=ttd_asn_values if ttd_asn_values else ["-"])
        self.combo_ttd_asn.set(ttd_asn_values[0] if ttd_asn_values else "-")

    def calculate_duration(self, *args):
        try:
            if HAS_TKCALENDAR:
                s_date, e_date = self.dp_mulai.get_date(), self.dp_akhir.get_date()
            else:
                s_date = datetime.strptime(self.dp_mulai.get(), "%d/%m/%Y")
                e_date = datetime.strptime(self.dp_akhir.get(), "%d/%m/%Y")
            delta = (e_date - s_date).days + 1
            if delta < 1: delta = 1
            self.ent_lama_hari.configure(state="normal")
            self.ent_lama_hari.delete(0, tk.END)
            self.ent_lama_hari.insert(0, str(delta))
            self.ent_lama_hari.configure(state="readonly")
        except:
            pass
        self.schedule_preview_refresh()

    # ------------------------------------------------------------------
    # CONTEXT BUILDING
    # ------------------------------------------------------------------
    def build_context(self, record_history=True):
        try:
            jml_angka = int(self.ent_lama_hari.get())
            jml_teks = self.terbilang(jml_angka)
        except:
            jml_angka = self.ent_lama_hari.get()
            jml_teks = "-"

        ttd_dprd_raw = self.combo_ttd_dprd.get()
        ttd_asn_raw = self.combo_ttd_asn.get()
        jab_dprd, nama_dprd = ttd_dprd_raw.split(" - ", 1) if " - " in ttd_dprd_raw else ("KETUA", ttd_dprd_raw)
        jab_asn, nama_asn = ttd_asn_raw.split(" - ", 1) if " - " in ttd_asn_raw else ("SEKRETARIS DPRD", ttd_asn_raw)

        if self.mode == "dprd":
            selected_dprd = [p for p in self.db_dprd
                             if self.dprd_vars.get(p.get('nama', '')) and self.dprd_vars[p.get('nama', '')].get()]
            selected_asn = [p for p in self.db_asn
                            if self.asn_vars.get(p.get('nama', '')) and self.asn_vars[p.get('nama', '')].get()]
            selected_pelaksana = []
            selected_pendamping = []
        else:
            selected_dprd = []
            selected_asn = []
            selected_pelaksana = [p for p in self.db_asn
                                 if self.pelaksana_vars.get(f"pelaksana_{p.get('nama', '')}") and
                                    self.pelaksana_vars[f"pelaksana_{p.get('nama', '')}"].get()]
            selected_pendamping = [p for p in self.db_asn
                                   if self.pendamping_vars.get(f"pendamping_{p.get('nama', '')}") and
                                      self.pendamping_vars[f"pendamping_{p.get('nama', '')}"].get()]

        if HAS_TKCALENDAR:
            tanggal_surat = self.format_indonesian_date(self.dp_surat.get_date())
            tanggal_mulai = self.format_indonesian_date(self.dp_mulai.get_date())
            tanggal_akhir = self.format_indonesian_date(self.dp_akhir.get_date())
        else:
            tanggal_surat = self.dp_surat.get()
            tanggal_mulai = self.dp_mulai.get()
            tanggal_akhir = self.dp_akhir.get()

        tujuan_list = self.tujuan_terpilih if self.tujuan_terpilih else ["(belum diisi)"]
        tujuan_str = " / ".join(tujuan_list)

        jenis_perjalanan = self.combo_jenis.get()
        materi_st = self.txt_materi_st.get("1.0", tk.END).strip()
        materi_pb = self.txt_materi_pb.get("1.0", tk.END).strip()
        dasar_surat_dprd = self.txt_dasar_dprd.get("1.0", tk.END).strip()
        dasar_surat_asn = self.txt_dasar_asn.get("1.0", tk.END).strip()

        nomor_spd_dprd = self.inputs["nomor_spd_dprd"].get()
        nomor_spd_asn = self.inputs["nomor_spd_asn"].get()
        nomor_spd_pelaksana = self.ent_spd_pelaksana.get() if self.mode == "setwan" else ""
        nomor_spd_pendamping = self.ent_spd_pendamping.get() if self.mode == "setwan" else ""

        city_names = [extract_city_name(d) for d in tujuan_list]
        if any(not is_in_sulawesi_utara(c) for c in city_names):
            transport = "Pesawat / Mobil / Kereta"
        else:
            transport = "Mobil"

        ctx = {
            "nomor_surat": self.inputs["nomor_surat"].get(),
            "nomor_surat_asn": self.inputs["nomor_surat_asn"].get(),
            "nomor_pemberitahuan_dprd": self.inputs["nomor_pemberitahuan_dprd"].get(),
            "nomor_pemberitahuan_asn": self.inputs["nomor_pemberitahuan_asn"].get(),
            "nomor_spd_dprd": nomor_spd_dprd,
            "nomor_spd_asn": nomor_spd_asn,
            "nomor_spd_pelaksana": nomor_spd_pelaksana,
            "nomor_spd_pendamping": nomor_spd_pendamping,
            "nomor_spd": nomor_spd_dprd,
            "tanggal_surat": tanggal_surat,
            "tanggal_surat_asn": tanggal_surat,
            "jenis_perjalanan": jenis_perjalanan,
            "tujuan_bertugas": tujuan_str,
            "tujuan_bertugas_list": tujuan_list,
            "dasar_surat_dprd": dasar_surat_dprd,
            "dasar_surat_asn": dasar_surat_asn,
            "materi_tugas": materi_st,
            "materi_tugas_asn": materi_st,
            "isi_surat_pemberitahuan": materi_pb,
            "isi_surat_izin": materi_st,
            "tanggal_mulai": tanggal_mulai,
            "tanggal_akhir": tanggal_akhir,
            "jumlah_angka": jml_angka,
            "jumlah_teks": jml_teks,
            "jabatan_ttd": jab_dprd.strip(),
            "nama_ttd": nama_dprd.strip(),
            "jabatan_ttd_asn": jab_asn.strip(),
            "nama_ttd_asn": nama_asn.strip(),
            "transportasi_otomatis": transport,
            "tanggal_surat_info": tanggal_surat,
            "tujuan_surat_info": tujuan_str,
            "pelaksana_tugas_dprd_info": "Pimpinan dan Anggota",
            "jenis_perjalanan_info": jenis_perjalanan,
            "tujuan_bertugas_info": tujuan_str,
            "materi_tugas_info": materi_pb,
            "hari_info": "Sesuai Jadwal",
            "tanggal_bertugas_info": f"{tanggal_mulai} s/d {tanggal_akhir}",
            "pelaksana_tugas_info": "Anggota DPRD",
            "jlh_pelaksana_dprd": len(selected_dprd),
            "pelaksana_tugas_asn_info": "Pendamping ASN" if self.mode == "dprd" else "Pelaksana ASN",
            "jlh_pelaksana_asn": len(selected_asn),
            "jlh_pelaksana": len(selected_pelaksana),
            "jlh_pendamping": len(selected_pendamping),
            "jabatan_ttd_info": jab_dprd.strip(),
            "nama_ttd_info": nama_dprd.strip(),
            "pelaksana_dprd": selected_dprd,
            "pelaksana_asn": selected_asn,
            "pelaksana_list": selected_pelaksana,
            "pendamping_list": selected_pendamping,
        }

        if record_history:
            self.history_data[ctx["nomor_surat"]] = {
                "nomor_surat": ctx["nomor_surat"],
                "nomor_surat_asn": ctx["nomor_surat_asn"],
                "nomor_pemberitahuan_dprd": ctx["nomor_pemberitahuan_dprd"],
                "nomor_pemberitahuan_asn": ctx["nomor_pemberitahuan_asn"],
                "nomor_spd_dprd": nomor_spd_dprd,
                "nomor_spd_asn": nomor_spd_asn,
                "jenis_perjalanan": ctx["jenis_perjalanan"],
                "tujuan_bertugas": tujuan_str,
                "tujuan_bertugas_list": tujuan_list,
                "materi_tugas": materi_st,
                "materi_tugas_pb": materi_pb,
                "dasar_surat_dprd": dasar_surat_dprd,
                "dasar_surat_asn": dasar_surat_asn,
                "dprd_terpilih": [p.get('nama') for p in selected_dprd],
                "asn_terpilih": [p.get('nama') for p in selected_asn],
                "pelaksana_terpilih": [p.get('nama') for p in selected_pelaksana] if self.mode == "setwan" else [],
                "pendamping_terpilih": [p.get('nama') for p in selected_pendamping] if self.mode == "setwan" else [],
            }
            self.save_history()
            if hasattr(self, "combo_history"):
                self.combo_history.configure(values=list(self.history_data.keys()))

        return ctx, selected_dprd, selected_asn, selected_pelaksana, selected_pendamping

    # ------------------------------------------------------------------
    # CETAK DOKUMEN
    # ------------------------------------------------------------------
    def generate_documents_action(self):
        ctx, sel_dprd, sel_asn, sel_pelaksana, sel_pendamping = self.build_context()
        if self.mode == "dprd":
            if not sel_dprd and not sel_asn:
                messagebox.showwarning("Peringatan", "Pilih minimal satu pelaksana!")
                return
        else:
            if not sel_pelaksana and not sel_pendamping:
                messagebox.showwarning("Peringatan", "Pilih minimal satu Pelaksana ASN atau Pendamping ASN!")
                return
        if not self.tujuan_terpilih:
            messagebox.showwarning("Peringatan", "Tambahkan minimal satu kota tujuan bertugas!")
            return

        out_dir = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan Dokumen")
        if not out_dir: return

        success_count = 0
        missing = []

        if self.mode == "dprd":
            if sel_dprd:
                template_st_dprd = TEMPLATE_ST_DPRD_BIASA if len(sel_dprd) <= 3 else TEMPLATE_ST_DPRD_TABEL
                if os.path.exists(template_st_dprd):
                    try:
                        out_path = os.path.join(out_dir, "Surat_Tugas_DPRD.docx")
                        buat_surat_tugas_dprd(ctx, sel_dprd, out_path)
                        success_count += 1
                    except Exception as e:
                        print(f"Gagal buat surat tugas DPRD: {e}")
                else:
                    missing.append(template_st_dprd)
            if sel_asn:
                template_st_asn = TEMPLATE_ST_ASN_BIASA if len(sel_asn) <= 3 else TEMPLATE_ST_ASN_TABEL
                if os.path.exists(template_st_asn):
                    try:
                        out_path = os.path.join(out_dir, "Surat_Tugas_ASN.docx")
                        buat_surat_tugas_asn(ctx, sel_asn, out_path)
                        success_count += 1
                    except Exception as e:
                        print(f"Gagal buat surat tugas ASN: {e}")
                else:
                    missing.append(template_st_asn)
            if os.path.exists("pemberitahuan_dprd.docx"):
                try:
                    base_number = ctx.get("nomor_pemberitahuan_dprd", ctx.get("nomor_surat", ""))
                    out_path = os.path.join(out_dir, "Surat_Pemberitahuan.docx")
                    buat_surat_pemberitahuan_multi(
                        "pemberitahuan_dprd.docx",
                        ctx,
                        sel_dprd,
                        sel_asn,
                        self.tujuan_terpilih,
                        base_number,
                        out_path,
                        label_asn="Pendamping ASN"
                    )
                    success_count += 1
                except Exception as e:
                    print(f"Gagal buat pemberitahuan: {e}")
            else:
                missing.append("pemberitahuan_dprd.docx")
            if sel_dprd:
                out_depan_dprd = os.path.join(out_dir, "SPD_DPRD_Depan.docx")
                out_belakang_dprd = os.path.join(out_dir, "SPD_DPRD_Belakang.docx")
                try:
                    buat_sppd_dprd("SPD_DPRD.docx", "SPD_BELAKANG.docx", ctx, sel_dprd,
                                    self.tujuan_terpilih, out_depan_dprd, out_belakang_dprd)
                    success_count += 2
                except Exception as e:
                    print(f"Gagal buat SPD DPRD: {e}")
            if sel_asn:
                out_depan_asn = os.path.join(out_dir, "SPD_ASN_Depan.docx")
                out_belakang_asn = os.path.join(out_dir, "SPD_ASN_Belakang.docx")
                try:
                    buat_sppd_asn("SPD_DPRD.docx", "SPD_BELAKANG.docx", ctx, sel_asn,
                                   self.tujuan_terpilih, out_depan_asn, out_belakang_asn)
                    success_count += 2
                except Exception as e:
                    print(f"Gagal buat SPD ASN: {e}")
        else:
            sel_asn = sel_pelaksana + sel_pendamping
            if sel_pelaksana:
                template_st_asn = TEMPLATE_ST_ASN_BIASA if len(sel_pelaksana) <= 3 else TEMPLATE_ST_ASN_TABEL
                if os.path.exists(template_st_asn):
                    try:
                        out_path = os.path.join(out_dir, "Surat_Tugas_Pelaksana_ASN.docx")
                        ctx_pelaksana = ctx.copy()
                        ctx_pelaksana["pelaksana_tugas_asn_info"] = "Pelaksana ASN"
                        ctx_pelaksana["jlh_pelaksana_asn"] = len(sel_pelaksana)
                        buat_surat_tugas_asn(ctx_pelaksana, sel_pelaksana, out_path)
                        success_count += 1
                    except Exception as e:
                        print(f"Gagal buat surat tugas Pelaksana ASN: {e}")
                else:
                    missing.append(template_st_asn)
            if sel_pendamping:
                template_st_asn = TEMPLATE_ST_ASN_BIASA if len(sel_pendamping) <= 3 else TEMPLATE_ST_ASN_TABEL
                if os.path.exists(template_st_asn):
                    try:
                        out_path = os.path.join(out_dir, "Surat_Tugas_Pendamping_ASN.docx")
                        ctx_pendamping = ctx.copy()
                        ctx_pendamping["pelaksana_tugas_asn_info"] = "Pendamping ASN"
                        ctx_pendamping["jlh_pelaksana_asn"] = len(sel_pendamping)
                        buat_surat_tugas_asn(ctx_pendamping, sel_pendamping, out_path)
                        success_count += 1
                    except Exception as e:
                        print(f"Gagal buat surat tugas Pendamping ASN: {e}")
                else:
                    missing.append(template_st_asn)
            if os.path.exists("pemberitahuan_dprd.docx"):
                if sel_pelaksana:
                    try:
                        base_number = ctx.get("nomor_pemberitahuan_asn", ctx.get("nomor_surat_asn", ""))
                        out_path = os.path.join(out_dir, "Surat_Pemberitahuan_Pelaksana_ASN.docx")
                        buat_surat_pemberitahuan_multi(
                            "pemberitahuan_dprd.docx",
                            ctx,
                            [],
                            sel_pelaksana,
                            self.tujuan_terpilih,
                            base_number,
                            out_path,
                            label_asn="Pelaksana ASN"
                        )
                        success_count += 1
                    except Exception as e:
                        print(f"Gagal buat pemberitahuan Pelaksana: {e}")
                if sel_pendamping:
                    try:
                        base_number = ctx.get("nomor_pemberitahuan_asn", ctx.get("nomor_surat_asn", ""))
                        out_path = os.path.join(out_dir, "Surat_Pemberitahuan_Pendamping_ASN.docx")
                        buat_surat_pemberitahuan_multi(
                            "pemberitahuan_dprd.docx",
                            ctx,
                            [],
                            sel_pendamping,
                            self.tujuan_terpilih,
                            base_number,
                            out_path,
                            label_asn="Pendamping ASN"
                        )
                        success_count += 1
                    except Exception as e:
                        print(f"Gagal buat pemberitahuan Pendamping: {e}")
            else:
                missing.append("pemberitahuan_dprd.docx")
            if sel_pelaksana:
                ctx_pelaksana_spd = ctx.copy()
                ctx_pelaksana_spd["nomor_spd_asn"] = ctx.get("nomor_spd_pelaksana", ctx.get("nomor_spd_asn", ""))
                out_depan_pelaksana = os.path.join(out_dir, "SPD_Pelaksana_ASN_Depan.docx")
                out_belakang_pelaksana = os.path.join(out_dir, "SPD_Pelaksana_ASN_Belakang.docx")
                try:
                    buat_sppd_asn("SPD_DPRD.docx", "SPD_BELAKANG.docx", ctx_pelaksana_spd, sel_pelaksana,
                                   self.tujuan_terpilih, out_depan_pelaksana, out_belakang_pelaksana)
                    success_count += 2
                except Exception as e:
                    print(f"Gagal buat SPD Pelaksana ASN: {e}")
            if sel_pendamping:
                ctx_pendamping_spd = ctx.copy()
                ctx_pendamping_spd["nomor_spd_asn"] = ctx.get("nomor_spd_pendamping", ctx.get("nomor_spd_asn", ""))
                out_depan_pendamping = os.path.join(out_dir, "SPD_Pendamping_ASN_Depan.docx")
                out_belakang_pendamping = os.path.join(out_dir, "SPD_Pendamping_ASN_Belakang.docx")
                try:
                    buat_sppd_asn("SPD_DPRD.docx", "SPD_BELAKANG.docx", ctx_pendamping_spd, sel_pendamping,
                                   self.tujuan_terpilih, out_depan_pendamping, out_belakang_pendamping)
                    success_count += 2
                except Exception as e:
                    print(f"Gagal buat SPD Pendamping ASN: {e}")

        # DAFTAR HADIR
        try:
            out_daftar = os.path.join(out_dir, "Daftar_Hadir.docx")
            if self.mode == "dprd":
                pelaksana_daftar = sel_dprd
            else:
                pelaksana_daftar = sel_pelaksana + sel_pendamping
            if pelaksana_daftar:
                self.buat_daftar_hadir(ctx, pelaksana_daftar, self.tujuan_terpilih, self.mode, out_daftar)
                success_count += 1
        except Exception as e:
            print(f"Gagal buat daftar hadir: {e}")

        self.combo_history.configure(values=list(self.history_data.keys()))

        if success_count > 0:
            info_lines = [f"Berhasil mencetak {success_count} file dokumen di:\n{out_dir}\n"]
            if self.mode == "dprd":
                if sel_dprd:
                    info_lines.append(f"SPD DPRD: {len(sel_dprd)} orang → nomor SPD: {ctx['nomor_spd_dprd']} (sama semua)")
                if sel_asn:
                    nomor_terakhir = increment_nomor_spd(ctx['nomor_spd_asn'], len(sel_asn) - 1)
                    info_lines.append(f"SPD ASN: {len(sel_asn)} orang → nomor: {ctx['nomor_spd_asn']} s/d {nomor_terakhir}")
            else:
                if sel_pelaksana:
                    nomor_spd_pelaksana = ctx.get('nomor_spd_pelaksana', '')
                    nomor_terakhir = increment_nomor_spd(nomor_spd_pelaksana, len(sel_pelaksana) - 1)
                    info_lines.append(f"Pelaksana ASN: {len(sel_pelaksana)} orang → SPD: {nomor_spd_pelaksana} s/d {nomor_terakhir}")
                if sel_pendamping:
                    nomor_spd_pendamping = ctx.get('nomor_spd_pendamping', '')
                    nomor_terakhir = increment_nomor_spd(nomor_spd_pendamping, len(sel_pendamping) - 1)
                    info_lines.append(f"Pendamping ASN: {len(sel_pendamping)} orang → SPD: {nomor_spd_pendamping} s/d {nomor_terakhir}")
            info_lines.append(f"Tujuan: {' / '.join(self.tujuan_terpilih)}")
            if missing:
                info_lines.append("\nTemplate tidak ditemukan:\n" + "\n".join(missing))
            messagebox.showinfo("Cetak Berhasil", "\n".join(info_lines))
        else:
            messagebox.showerror("Error", "Gagal mencetak. Pastikan file template tersedia dan ada pelaksana terpilih.")

    # ------------------------------------------------------------------
    # DAFTAR HADIR
    # ------------------------------------------------------------------
    def buat_daftar_hadir(self, ctx, pelaksana_list, destinations, mode, out_path):
        template_path = "DAFTAR_HADIR_DPRD.docx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template {template_path} tidak ditemukan.")

        periods = generate_periods(ctx.get("tanggal_mulai", ""), destinations)

        temp_files = []
        for period in periods:
            tujuan = period["tujuan"]
            hari = period["hari"]
            tanggal = period["tanggal"]

            judul_pelaksana, judul_tujuan = self.build_judul_daftar_hadir(
                pelaksana_list, tujuan, ctx.get("jenis_perjalanan", ""), ctx.get("materi_tugas", ""), mode
            )

            render_ctx = {
                "MATERI_TUGAS_DPRD_DAFTAR_HADIR": judul_pelaksana,
                "TEMPAT_TUGAS_DPRD_DAFTAR_HADIR": judul_tujuan,
                "HARI": hari,
                "TANGGAL_DAFTAR_HADIR": tanggal,
                "TEMPAT_DAFTAR_HADIR": tujuan,
                "loop": {"index": ""},
                "tabel": {"NAMA_DAFTAR_HADIR": "", "jabatan_daftar_hadir": ""}
            }

            doc_tpl = DocxTemplate(template_path)
            doc_tpl.render(render_ctx)
            tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
            doc_tpl.save(tmp_file)

            doc = Document(tmp_file)
            header_keywords = ["no", "nama", "jabatan", "tanda tangan"]
            rows_data = []
            for i, p in enumerate(pelaksana_list):
                rows_data.append([
                    str(i + 1),
                    p.get('nama', ''),
                    p.get('jabatan', '')
                ])
            _fill_table_rows_from_master(doc, header_keywords, rows_data)
            doc.save(tmp_file)
            temp_files.append(tmp_file)

        if temp_files:
            _combine_word_pages(temp_files, out_path)
        else:
            raise Exception("Tidak ada halaman yang dihasilkan.")

        for f in temp_files:
            try:
                os.unlink(f)
            except:
                pass

    def build_judul_daftar_hadir(self, pelaksana_list, tujuan, jenis_perjalanan, materi, mode):
        kategori_set = set()
        for p in pelaksana_list:
            kat = p.get('kategori', '').strip()
            if kat:
                kategori_set.add(kat)

        if len(kategori_set) == 1:
            kategori = list(kategori_set)[0]
            if kategori in ["Pimpinan DPRD", "Komisi I", "Komisi II", "Komisi III"]:
                has_pimpinan = any(
                    "ketua" in p.get('jabatan', '').lower() or
                    "wakil" in p.get('jabatan', '').lower() or
                    "sekretaris" in p.get('jabatan', '').lower()
                    for p in pelaksana_list
                )
                has_anggota = any(
                    "anggota" in p.get('jabatan', '').lower()
                    for p in pelaksana_list
                )
                if has_pimpinan and has_anggota:
                    label = f"Pimpinan dan Anggota {kategori}"
                elif has_pimpinan:
                    label = f"Pimpinan {kategori}"
                elif has_anggota:
                    label = f"Anggota {kategori}"
                else:
                    label = kategori
                if "DPRD" not in label:
                    label = f"{label} DPRD Kota Bitung"
                pelaku_str = label.upper()
            else:
                pelaku_str = "PIMPINAN DAN ANGGOTA DPRD KOTA BITUNG"
        else:
            pelaku_str = "PIMPINAN DAN ANGGOTA DPRD KOTA BITUNG"

        tujuan_daerah = extract_city_name(tujuan)
        if "DPRD" not in tujuan_daerah.upper():
            instansi_tujuan = f"DPRD {tujuan_daerah}"
        else:
            instansi_tujuan = tujuan_daerah

        jenis = jenis_perjalanan.upper().strip()
        materi_upper = materi.upper().strip()

        judul_pelaksana = f"{pelaku_str} PADA {jenis} KE {instansi_tujuan} TENTANG {materi_upper}"
        judul_tujuan = f"{instansi_tujuan} PADA {jenis} {pelaku_str} KE {instansi_tujuan} TENTANG {materi_upper}"

        return judul_pelaksana, judul_tujuan

    # ------------------------------------------------------------------
    # LIVE PREVIEW
    # ------------------------------------------------------------------
    def schedule_preview_refresh(self, *args, immediate=False, **kwargs):
        if not hasattr(self, "combo_preview_jenis"):
            return
        if self._preview_after_id is not None:
            try:
                self.after_cancel(self._preview_after_id)
            except Exception:
                pass
            self._preview_after_id = None
        delay = 100 if immediate else 700
        self._preview_after_id = self.after(delay, self._launch_preview_render)

    def _launch_preview_render(self):
        self._preview_after_id = None
        start_thread = False
        with self._preview_lock:
            if self._preview_busy:
                self._preview_pending = True
            else:
                self._preview_busy = True
                start_thread = True
        if start_thread:
            threading.Thread(target=self._preview_worker, daemon=True).start()

    def _set_preview_status(self, text, color="gray"):
        def _do():
            try:
                self.preview_status_lbl.configure(text=text, text_color=color)
            except Exception:
                pass
        self.after(0, _do)

    def _preview_worker(self):
        try:
            label = self.combo_preview_jenis.get()
            template_file, mode = None, "ctx"
            for name, fname, m in PREVIEW_TEMPLATES:
                if name == label:
                    template_file, mode = fname, m
                    break

            if not template_file:
                self._set_preview_status("Pilih jenis surat untuk pratinjau.")
                return
            if template_file not in ("__surat_tugas_dprd__", "__surat_tugas_asn__", "__daftar_hadir__") and not os.path.exists(template_file):
                self._set_preview_status(f"Template tidak ditemukan: {template_file}")
                return

            if self.current_view != "perjalanan_dinas":
                self._set_preview_status("Preview tidak tersedia di mode Surat Undangan.")
                return

            self._set_preview_status("Memproses preview...")
            ctx, sel_dprd, sel_asn, sel_pelaksana, sel_pendamping = self.build_context(record_history=False)

            tmp_docx = os.path.join(self.preview_dir, "preview_render.docx")
            tmp_pdf = os.path.join(self.preview_dir, "preview_render.pdf")

            if mode == "person_dprd":
                if not sel_dprd:
                    self._set_preview_status("Pilih minimal satu anggota DPRD untuk melihat pratinjau SPD DPRD.")
                    return
                nomor_dprd = ctx.get('nomor_spd_dprd', '')
                transport = ctx.get("transportasi_otomatis", "Mobil")
                render_ctx = _build_person_sppd_context(ctx, sel_dprd[0], nomor_dprd, self.tujuan_terpilih, transport)
                doc = DocxTemplate(template_file)
                doc.render(render_ctx)
                doc.save(tmp_docx)
            elif mode == "person_asn":
                if not sel_asn:
                    self._set_preview_status("Pilih minimal satu ASN untuk melihat pratinjau SPD ASN.")
                    return
                nomor_asn = ctx.get('nomor_spd_asn', '')
                transport = ctx.get("transportasi_otomatis", "Mobil")
                render_ctx = _build_person_sppd_context(ctx, sel_asn[0], nomor_asn, self.tujuan_terpilih, transport)
                doc = DocxTemplate(template_file)
                doc.render(render_ctx)
                doc.save(tmp_docx)
            elif template_file == "__surat_tugas_dprd__":
                if not sel_dprd:
                    self._set_preview_status("Pilih minimal satu pelaksana DPRD untuk pratinjau Surat Tugas DPRD.")
                    return
                buat_surat_tugas_dprd(ctx, sel_dprd, tmp_docx)
            elif template_file == "__surat_tugas_asn__":
                if not sel_asn:
                    self._set_preview_status("Pilih minimal satu pendamping ASN untuk pratinjau Surat Tugas ASN.")
                    return
                buat_surat_tugas_asn(ctx, sel_asn, tmp_docx)
            elif template_file == "pemberitahuan_dprd.docx":
                if not self.tujuan_terpilih:
                    self._set_preview_status("Tambahkan minimal satu tujuan untuk pratinjau pemberitahuan.")
                    return
                base_number = ctx.get("nomor_pemberitahuan_dprd", ctx.get("nomor_surat", ""))
                if self.mode == "dprd":
                    pelaksana_dprd = sel_dprd
                    pelaksana_asn = sel_asn
                    label_asn = "Pendamping ASN"
                else:
                    pelaksana_dprd = []
                    pelaksana_asn = sel_asn
                    label_asn = "Pelaksana Tugas ASN"
                buat_surat_pemberitahuan_multi(
                    template_file,
                    ctx,
                    pelaksana_dprd,
                    pelaksana_asn,
                    self.tujuan_terpilih[:1],
                    base_number,
                    tmp_docx,
                    label_asn=label_asn
                )
            elif template_file == "__daftar_hadir__":
                if self.mode == "dprd":
                    pelaksana = sel_dprd
                else:
                    pelaksana = sel_asn
                if not pelaksana:
                    self._set_preview_status("Pilih minimal satu pelaksana untuk pratinjau daftar hadir.")
                    return
                if not self.tujuan_terpilih:
                    self._set_preview_status("Tambahkan minimal satu tujuan untuk pratinjau daftar hadir.")
                    return
                try:
                    self.buat_daftar_hadir(ctx, pelaksana, self.tujuan_terpilih[:1], self.mode, tmp_docx)
                except Exception as e:
                    self._set_preview_status(f"Gagal membuat daftar hadir: {e}")
                    return
            else:
                doc = DocxTemplate(template_file)
                doc.render(ctx)
                doc.save(tmp_docx)

            if not HAS_FITZ:
                self._set_preview_status("Pratinjau visual butuh paket 'PyMuPDF'.\nJalankan: pip install pymupdf pillow")
                return

            if os.path.exists(tmp_pdf):
                try: os.remove(tmp_pdf)
                except Exception: pass

            soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
            if soffice_bin:
                ok = self._convert_with_soffice(soffice_bin, tmp_docx, tmp_pdf)
            else:
                ok = False

            if not ok and HAS_DOCX2PDF:
                ok = self._convert_with_docx2pdf_safe(tmp_docx, tmp_pdf)

            if not ok:
                self._set_preview_status(
                    "Gagal membuat PDF untuk pratinjau.\n"
                    "Pasang LibreOffice atau Microsoft Word di komputer ini."
                )
                return

            pdf_doc = fitz.open(tmp_pdf)
            page = pdf_doc.load_page(0)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img_mode = "RGBA" if pix.alpha else "RGB"
            img = Image.frombytes(img_mode, [pix.width, pix.height], pix.samples)
            pdf_doc.close()

            self.after(0, lambda: self._apply_preview_image(img))
        except Exception as e:
            self._set_preview_status(f"Gagal membuat preview: {str(e)}")
        finally:
            with self._preview_lock:
                self._preview_busy = False
                need_again = self._preview_pending
                self._preview_pending = False
            if need_again:
                self.after(50, self._launch_preview_render)

    def _apply_preview_image(self, pil_img):
        try:
            self.preview_canvas_frame.update_idletasks()
            frame_width = self.preview_canvas_frame.winfo_width()
            if frame_width < 100: frame_width = 480
        except Exception:
            frame_width = 480

        avail = max(frame_width - 30, 200)
        scale = avail / pil_img.width
        scale = max(0.12, min(scale, 1.2))
        disp_w = max(int(pil_img.width * scale), 50)
        disp_h = max(int(pil_img.height * scale), 50)

        ctk_img = ctk.CTkImage(light_image=pil_img, size=(disp_w, disp_h))
        self._preview_ctk_image = ctk_img
        self.preview_image_label.configure(image=ctk_img, text="")
        self.preview_status_lbl.configure(text="Pratinjau terkini.", text_color="gray")

    def _convert_with_soffice(self, soffice_bin, docx_path, pdf_path):
        out_dir = os.path.dirname(os.path.abspath(pdf_path)) or "."
        try:
            subprocess.run(
                [soffice_bin, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                check=True, timeout=60, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )
            generated = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if os.path.exists(generated):
                if os.path.abspath(generated) != os.path.abspath(pdf_path):
                    if os.path.exists(pdf_path): os.remove(pdf_path)
                    os.replace(generated, pdf_path)
                return True
            return False
        except Exception as e:
            print(f"Konversi LibreOffice gagal: {e}")
            return False

    def _convert_with_docx2pdf_safe(self, docx_path, pdf_path):
        com_ready = False
        try:
            import pythoncom
            pythoncom.CoInitialize()
            com_ready = True
        except Exception:
            pass
        try:
            tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
            shutil.copy2(docx_path, tmp_docx)
            tmp_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
            convert_to_pdf_word(tmp_docx, tmp_pdf)
            if os.path.exists(tmp_pdf):
                if os.path.exists(pdf_path): os.remove(pdf_path)
                os.replace(tmp_pdf, pdf_path)
                return True
            return False
        except Exception as e:
            print(f"Konversi docx2pdf gagal: {e}")
            return False
        finally:
            try:
                if os.path.exists(tmp_docx): os.unlink(tmp_docx)
                if os.path.exists(tmp_pdf): os.unlink(tmp_pdf)
            except:
                pass
            if com_ready:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    # ------------------------------------------------------------------
    # SURAT UNDANGAN
    # ------------------------------------------------------------------
    def _switch_to_undangan_layout(self):
        self.right_frame.grid_remove()
        self.preview_frame.grid(row=0, column=2, padx=10, pady=10, sticky="nsew")
        self.grid_columnconfigure(2, weight=2, minsize=480)
        self.grid_columnconfigure(3, weight=0, minsize=0)

    def _switch_to_perjalanan_layout(self):
        self.right_frame.grid(row=0, column=2, padx=10, pady=10, sticky="nsew")
        self.preview_frame.grid(row=0, column=3, padx=10, pady=10, sticky="nsew")
        self.grid_columnconfigure(2, weight=1, minsize=380)
        self.grid_columnconfigure(3, weight=2, minsize=480)

    def clear_preview(self):
        self.preview_image_label.configure(image=None, text="")

    def show_undangan_paripurna(self):
        self.current_view = "undangan_paripurna"
        self.middle_frame.configure(label_text="Surat Undangan Paripurna")
        self.btn_generate_main.configure(text="⚡ CETAK SURAT UNDANGAN", command=self.generate_undangan_paripurna)
        self.btn_back_to_perjalanan.grid()
        self.btn_undangan_paripurna.configure(state="disabled")
        self.btn_undangan_biasa.configure(state="normal")
        self._switch_to_undangan_layout()
        self.setup_undangan_paripurna_form()
        self.clear_preview()

    def show_undangan_biasa(self):
        self.current_view = "undangan_biasa"
        self.middle_frame.configure(label_text="Surat Undangan Biasa")
        self.btn_generate_main.configure(text="⚡ CETAK SURAT UNDANGAN", command=self.generate_undangan_biasa)
        self.btn_back_to_perjalanan.grid()
        self.btn_undangan_biasa.configure(state="disabled")
        self.btn_undangan_paripurna.configure(state="normal")
        self._switch_to_undangan_layout()
        messagebox.showinfo("Info", "Fitur Undangan Biasa akan segera tersedia.")

    def show_perjalanan_dinas(self):
        self.current_view = "perjalanan_dinas"
        self.middle_frame.configure(label_text="Data Perjalanan Dinas")
        self.btn_generate_main.configure(text="⚡ CETAK SURAT & SPD", command=self.generate_documents_action)
        self.btn_back_to_perjalanan.grid_remove()
        self.btn_undangan_paripurna.configure(state="normal")
        self.btn_undangan_biasa.configure(state="normal")
        self._switch_to_perjalanan_layout()
        self.setup_perjalanan_dinas_form()
        self.schedule_preview_refresh(immediate=True)

    def setup_undangan_paripurna_form(self):
        for widget in self.middle_frame.winfo_children():
            widget.destroy()
        self.undangan_inputs = {}

        # 1. Tanggal Surat
        lbl_tgl_surat = ctk.CTkLabel(self.middle_frame, text="1. Tanggal Surat", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_surat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.undangan_inputs["tanggal_surat"] = DateEntry(self.middle_frame, width=15, background='#2563EB',
                                                              foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.undangan_inputs["tanggal_surat"].pack(fill="x", padx=10, pady=(0, 6))
            self.undangan_inputs["tanggal_surat"].bind("<<DateEntrySelected>>", lambda e: self.schedule_preview_refresh())
        else:
            self.undangan_inputs["tanggal_surat"] = ctk.CTkEntry(self.middle_frame)
            self.undangan_inputs["tanggal_surat"].pack(fill="x", padx=10, pady=(0, 6))

        # 2. Nomor Undangan
        lbl_nomor = ctk.CTkLabel(self.middle_frame, text="2. Nomor Undangan", anchor="w", font=("Arial", 12, "bold"))
        lbl_nomor.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["nomor_undangan"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Contoh: 01/UNDP/X/2026")
        self.undangan_inputs["nomor_undangan"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["nomor_undangan"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        # 3. Isi Surat (Agenda)
        lbl_isi = ctk.CTkLabel(self.middle_frame, text="3. Isi Surat / Agenda Rapat", anchor="w", font=("Arial", 12, "bold"))
        lbl_isi.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["isi_surat"] = ctk.CTkTextbox(self.middle_frame, height=80, wrap="word")
        self.undangan_inputs["isi_surat"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["isi_surat"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        # 4. Tanggal Pelaksanaan Rapat
        lbl_tgl_rapat = ctk.CTkLabel(self.middle_frame, text="4. Tanggal Pelaksanaan Rapat", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_rapat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.undangan_inputs["tanggal_rapat"] = DateEntry(self.middle_frame, width=15, background='#2563EB',
                                                              foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.undangan_inputs["tanggal_rapat"].pack(fill="x", padx=10, pady=(0, 6))
            self.undangan_inputs["tanggal_rapat"].bind("<<DateEntrySelected>>", lambda e: self.update_hari_rapat())
        else:
            self.undangan_inputs["tanggal_rapat"] = ctk.CTkEntry(self.middle_frame)
            self.undangan_inputs["tanggal_rapat"].pack(fill="x", padx=10, pady=(0, 6))

        hari_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        hari_frame.pack(fill="x", padx=10, pady=(0, 6))
        lbl_hari = ctk.CTkLabel(hari_frame, text="Hari:", anchor="w", font=("Arial", 11))
        lbl_hari.pack(side="left", padx=(0, 10))
        self.undangan_inputs["hari_rapat"] = ctk.CTkEntry(hari_frame, width=150, state="readonly")
        self.undangan_inputs["hari_rapat"].pack(side="left")

        # 5. Jam Pelaksanaan
        lbl_jam = ctk.CTkLabel(self.middle_frame, text="5. Jam Pelaksanaan", anchor="w", font=("Arial", 12, "bold"))
        lbl_jam.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["jam_pelaksanaan"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Contoh: 09.00 WITA s.d. selesai")
        self.undangan_inputs["jam_pelaksanaan"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["jam_pelaksanaan"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        # 6. Skenario Rapat (Dynamic add/remove)
        lbl_skenario = ctk.CTkLabel(self.middle_frame, text="6. Skenario Rapat", anchor="w", font=("Arial", 12, "bold"))
        lbl_skenario.pack(fill="x", padx=10, pady=(8, 2))

        self.skenario_container = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        self.skenario_container.pack(fill="x", padx=10, pady=(0, 2))
        self.undangan_inputs["skenario"] = []

        self._add_skenario_row()

        btn_add_skenario = ctk.CTkButton(self.middle_frame, text="+ Tambah Skenario",
                                         command=self._add_skenario_row,
                                         fg_color="#6366F1", hover_color="#4F46E5", height=28)
        btn_add_skenario.pack(fill="x", padx=10, pady=(2, 6))

        # 7. Pakaian
        lbl_pakaian = ctk.CTkLabel(self.middle_frame, text="7. Pakaian", anchor="w", font=("Arial", 12, "bold"))
        lbl_pakaian.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["pakaian"] = ctk.CTkComboBox(self.middle_frame, values=["PSH", "PSR", "PSL"],
                                                          command=lambda choice: self.schedule_preview_refresh())
        self.undangan_inputs["pakaian"].pack(fill="x", padx=10, pady=(0, 6))
        self.undangan_inputs["pakaian"].set("PSH")

        # 8. Penandatanganan
        lbl_ttd = ctk.CTkLabel(self.middle_frame, text="8. Penandatanganan", anchor="w", font=("Arial", 12, "bold"))
        lbl_ttd.pack(fill="x", padx=10, pady=(8, 2))
        self.undangan_inputs["penandatanganan"] = ctk.CTkComboBox(self.middle_frame,
                                                                   values=[f"{p.get('jabatan', '')} - {p.get('nama', '')}" for p in self.db_dprd if "Pimpinan" in p.get("kategori", "")])
        self.undangan_inputs["penandatanganan"].pack(fill="x", padx=10, pady=(0, 6))

        self.update_hari_rapat()

    def setup_perjalanan_dinas_form(self):
        for widget in self.middle_frame.winfo_children():
            widget.destroy()
        self.inputs = {}
        self.mode_specific_widgets = {}

        for var_name, label in [
            ("nomor_surat", "Nomor Surat Tugas DPRD"),
            ("nomor_surat_asn", "Nomor Surat Tugas Setwan"),
        ]:
            lbl = ctk.CTkLabel(self.middle_frame, text=label, anchor="w", font=("Arial", 12, "bold"))
            lbl.pack(fill="x", padx=10, pady=(8, 2))
            ent = ctk.CTkEntry(self.middle_frame, placeholder_text=f"Masukkan {label.lower()}...")
            ent.pack(fill="x", padx=10, pady=(0, 6))
            ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
            self.inputs[var_name] = ent

        self.mode_specific_widgets["lbl_pemberitahuan_dprd"] = ctk.CTkLabel(self.middle_frame, text="Nomor Surat Pemberitahuan DPRD", anchor="w", font=("Arial", 12, "bold"))
        self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(8, 2))
        self.inputs["nomor_pemberitahuan_dprd"] = ctk.CTkEntry(self.middle_frame, placeholder_text="Masukkan nomor surat pemberitahuan dprd...")
        self.inputs["nomor_pemberitahuan_dprd"].pack(fill="x", padx=10, pady=(0, 6))
        self.inputs["nomor_pemberitahuan_dprd"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        self.mode_specific_widgets["ent_pemberitahuan_dprd"] = self.inputs["nomor_pemberitahuan_dprd"]

        for var_name, label in [
            ("nomor_pemberitahuan_asn", "Nomor Surat Pemberitahuan Setwan"),
        ]:
            lbl = ctk.CTkLabel(self.middle_frame, text=label, anchor="w", font=("Arial", 12, "bold"))
            lbl.pack(fill="x", padx=10, pady=(8, 2))
            ent = ctk.CTkEntry(self.middle_frame, placeholder_text=f"Masukkan {label.lower()}...")
            ent.pack(fill="x", padx=10, pady=(0, 6))
            ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
            self.inputs[var_name] = ent

        lbl_spd_title = ctk.CTkLabel(self.middle_frame, text="Nomor SPD", anchor="w", font=("Arial", 12, "bold"))
        lbl_spd_title.pack(fill="x", padx=10, pady=(10, 2))

        spd_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        spd_frame.pack(fill="x", padx=10, pady=(0, 6))
        spd_frame.grid_columnconfigure(0, weight=1)
        spd_frame.grid_columnconfigure(1, weight=1)

        self.mode_specific_widgets["lbl_spd_dprd"] = ctk.CTkLabel(spd_frame, text="Nomor SPD DPRD :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        self.mode_specific_widgets["lbl_spd_dprd"].grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.inputs["nomor_spd_dprd"] = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 10/SPD/X/2026/")
        self.inputs["nomor_spd_dprd"].grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.inputs["nomor_spd_dprd"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())
        self.mode_specific_widgets["ent_spd_dprd"] = self.inputs["nomor_spd_dprd"]

        self.mode_specific_widgets["lbl_spd_setwan"] = ctk.CTkLabel(spd_frame, text="Nomor SPD Setwan (ASN) :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        self.mode_specific_widgets["lbl_spd_setwan"].grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.inputs["nomor_spd_asn"] = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 20/SPD/X/2026/")
        self.inputs["nomor_spd_asn"].grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.inputs["nomor_spd_asn"].bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        self.lbl_spd_pelaksana = ctk.CTkLabel(spd_frame, text="Nomor SPD Pelaksana ASN :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        self.ent_spd_pelaksana = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 10/SPD-PL/X/2026/")
        self.ent_spd_pelaksana.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        self.lbl_spd_pendamping = ctk.CTkLabel(spd_frame, text="Nomor SPD Pendamping ASN :", anchor="w", font=("Arial", 11), text_color="#059669")
        self.ent_spd_pendamping = ctk.CTkEntry(spd_frame, placeholder_text="Contoh: 20/SPD-PD/X/2026/")
        self.ent_spd_pendamping.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_spd_info = ctk.CTkLabel(self.middle_frame, text="ℹ️  SPD DPRD: semua anggota pakai nomor sama  |  SPD ASN: nomor otomatis berurutan", anchor="w", font=("Arial", 10), text_color="gray")
        lbl_spd_info.pack(fill="x", padx=10, pady=(0, 6))
        self.mode_specific_widgets["lbl_spd_info"] = lbl_spd_info

        # Continue with rest of the form...
        self._build_remaining_perjalanan_form()

    def _build_remaining_perjalanan_form(self):
        lbl_dasar_title = ctk.CTkLabel(self.middle_frame, text="Dasar Surat Tugas", anchor="w", font=("Arial", 12, "bold"))
        lbl_dasar_title.pack(fill="x", padx=10, pady=(10, 2))

        dasar_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        dasar_frame.pack(fill="x", padx=10, pady=(0, 6))
        dasar_frame.grid_columnconfigure(0, weight=1)
        dasar_frame.grid_columnconfigure(1, weight=1)

        lbl_dasar_dprd = ctk.CTkLabel(dasar_frame, text="Dasar Surat Tugas DPRD :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        lbl_dasar_dprd.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.txt_dasar_dprd = ctk.CTkTextbox(dasar_frame, height=60, wrap="word")
        self.txt_dasar_dprd.grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.txt_dasar_dprd.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_dasar_asn = ctk.CTkLabel(dasar_frame, text="Dasar Surat Tugas ASN :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        lbl_dasar_asn.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.txt_dasar_asn = ctk.CTkTextbox(dasar_frame, height=60, wrap="word")
        self.txt_dasar_asn.grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.txt_dasar_asn.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_materi_title = ctk.CTkLabel(self.middle_frame, text="Materi / Agenda Kegiatan", anchor="w", font=("Arial", 12, "bold"))
        lbl_materi_title.pack(fill="x", padx=10, pady=(10, 2))

        materi_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        materi_frame.pack(fill="x", padx=10, pady=(0, 6))
        materi_frame.grid_columnconfigure(0, weight=1)
        materi_frame.grid_columnconfigure(1, weight=1)

        lbl_mt_st = ctk.CTkLabel(materi_frame, text="Surat Tugas & SPPD :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        lbl_mt_st.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
        self.txt_materi_st = ctk.CTkTextbox(materi_frame, height=70, wrap="word")
        self.txt_materi_st.grid(row=1, column=0, padx=(0, 4), sticky="ew")
        self.txt_materi_st.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_mt_pb = ctk.CTkLabel(materi_frame, text="Surat Pemberitahuan :", anchor="w", font=("Arial", 11), text_color="#1E3A8A")
        lbl_mt_pb.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
        self.txt_materi_pb = ctk.CTkTextbox(materi_frame, height=70, wrap="word")
        self.txt_materi_pb.grid(row=1, column=1, padx=(4, 0), sticky="ew")
        self.txt_materi_pb.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_jp = ctk.CTkLabel(self.middle_frame, text="Jenis Perjalanan", anchor="w", font=("Arial", 12, "bold"))
        lbl_jp.pack(fill="x", padx=10, pady=(8, 2))
        self.combo_jenis = ctk.CTkComboBox(self.middle_frame, values=["Kunjungan Kerja", "Kunjungan Konsultasi", "Bimbingan Teknis"], command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_jenis.pack(fill="x", padx=10, pady=(0, 6))

        lbl_tgl_surat = ctk.CTkLabel(self.middle_frame, text="Tanggal Surat", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_surat.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_surat = DateEntry(self.middle_frame, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_surat.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_surat.bind("<<DateEntrySelected>>", lambda e: self.schedule_preview_refresh())
        else:
            self.dp_surat = ctk.CTkEntry(self.middle_frame)
            self.dp_surat.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_surat.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_tgl_mulai = ctk.CTkLabel(self.middle_frame, text="Tanggal Mulai Tugas", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_mulai.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_mulai = DateEntry(self.middle_frame, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_mulai.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_mulai.bind("<<DateEntrySelected>>", lambda e: self.calculate_duration())
        else:
            self.dp_mulai = ctk.CTkEntry(self.middle_frame)
            self.dp_mulai.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_mulai.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_tgl_akhir = ctk.CTkLabel(self.middle_frame, text="Tanggal Selesai Tugas", anchor="w", font=("Arial", 12, "bold"))
        lbl_tgl_akhir.pack(fill="x", padx=10, pady=(8, 2))
        if HAS_TKCALENDAR:
            self.dp_akhir = DateEntry(self.middle_frame, width=15, background='#2563EB', foreground='white', borderwidth=1, date_pattern='dd/mm/yyyy')
            self.dp_akhir.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_akhir.bind("<<DateEntrySelected>>", lambda e: self.calculate_duration())
        else:
            self.dp_akhir = ctk.CTkEntry(self.middle_frame)
            self.dp_akhir.pack(fill="x", padx=10, pady=(0, 6))
            self.dp_akhir.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        lbl_lama = ctk.CTkLabel(self.middle_frame, text="Lama Tugas (Hari)", anchor="w", font=("Arial", 12, "bold"))
        lbl_lama.pack(fill="x", padx=10, pady=(8, 2))
        self.ent_lama_hari = ctk.CTkEntry(self.middle_frame, state="readonly")
        self.ent_lama_hari.pack(fill="x", padx=10, pady=(0, 6))

        lbl_tujuan = ctk.CTkLabel(self.middle_frame, text="Kota Tujuan Bertugas (Multi Lokasi)", anchor="w", font=("Arial", 12, "bold"))
        lbl_tujuan.pack(fill="x", padx=10, pady=(8, 2))
        tujuan_input_frame = ctk.CTkFrame(self.middle_frame, fg_color="transparent")
        tujuan_input_frame.pack(fill="x", padx=10, pady=(0, 2))
        tujuan_input_frame.grid_columnconfigure(0, weight=1)
        self.ent_tujuan = ctk.CTkEntry(tujuan_input_frame, placeholder_text="Ketik nama kota lalu klik Tambah...")
        self.ent_tujuan.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.ent_tujuan.bind("<KeyRelease>", self.on_tujuan_key_release)
        self.btn_tambah_tujuan = ctk.CTkButton(tujuan_input_frame, text="+ Tambah", width=80, command=self.tambah_tujuan)
        self.btn_tambah_tujuan.grid(row=0, column=1, sticky="e")

        self.suggestion_frame = ctk.CTkScrollableFrame(self.middle_frame, height=110, fg_color="#F3F4F6")

        lbl_tujuan_terpilih = ctk.CTkLabel(self.middle_frame, text="Lokasi yang dipilih:", anchor="w", font=("Arial", 11), text_color="gray")
        lbl_tujuan_terpilih.pack(fill="x", padx=10, pady=(4, 1))
        self.tujuan_list_frame = ctk.CTkScrollableFrame(self.middle_frame, height=80, fg_color="#F0F4FF")
        self.tujuan_list_frame.pack(fill="x", padx=10, pady=(0, 6))

        lbl_tujuan_hint = ctk.CTkLabel(self.middle_frame, text="ℹ️  Klik ✕ pada lokasi untuk menghapus. Urutan sesuai tampilan.", anchor="w", font=("Arial", 10), text_color="gray")
        lbl_tujuan_hint.pack(fill="x", padx=10, pady=(0, 6))

        self.lbl_sign_dprd = ctk.CTkLabel(self.middle_frame, text="Penandatangan DPRD:", font=("Arial", 12, "bold"))
        self.lbl_sign_dprd.pack(fill="x", padx=10, pady=(15, 2))
        self.combo_ttd_dprd = ctk.CTkComboBox(self.middle_frame, values=["-"], height=32,
                                               command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_ttd_dprd.pack(fill="x", padx=10, pady=(0, 8))

        self.lbl_sign_asn = ctk.CTkLabel(self.middle_frame, text="Penandatangan ASN / SPPD:", font=("Arial", 12, "bold"))
        self.lbl_sign_asn.pack(fill="x", padx=10, pady=(10, 2))
        self.combo_ttd_asn = ctk.CTkComboBox(self.middle_frame, values=["-"], height=32,
                                              command=lambda choice: self.schedule_preview_refresh(immediate=True))
        self.combo_ttd_asn.pack(fill="x", padx=10, pady=(0, 8))

        self.refresh_signer_dropdowns()
        self.refresh_tujuan_list_ui()

        if self.mode == "setwan":
            self.lbl_sign_dprd.pack_forget()
            self.combo_ttd_dprd.pack_forget()
            self.mode_specific_widgets["lbl_pemberitahuan_dprd"].pack_forget()
            self.mode_specific_widgets["ent_pemberitahuan_dprd"].pack_forget()
            self.mode_specific_widgets["lbl_spd_dprd"].grid_forget()
            self.mode_specific_widgets["ent_spd_dprd"].grid_forget()
            self.mode_specific_widgets["lbl_spd_setwan"].grid_forget()
            self.mode_specific_widgets["lbl_spd_info"].configure(text="ℹ️  SPD Pelaksana & Pendamping: nomor otomatis berurutan per kategori")
            self.lbl_spd_pelaksana.grid(row=0, column=0, padx=(0, 4), pady=(0, 2), sticky="w")
            self.ent_spd_pelaksana.grid(row=1, column=0, padx=(0, 4), sticky="ew")
            self.lbl_spd_pendamping.grid(row=0, column=1, padx=(4, 0), pady=(0, 2), sticky="w")
            self.ent_spd_pendamping.grid(row=1, column=1, padx=(4, 0), sticky="ew")
            self.combo_jenis.configure(values=["Studi Komparasi", "Kunjungan Konsultasi", "Bimbingan Teknis"])

    def _add_skenario_row(self):
        idx = len(self.undangan_inputs["skenario"]) + 1
        row_frame = ctk.CTkFrame(self.skenario_container, fg_color="transparent")
        row_frame.pack(fill="x", pady=2)
        row_frame.grid_columnconfigure(0, weight=1)

        ent = ctk.CTkEntry(row_frame, placeholder_text=f"Skenario {idx}...")
        ent.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        ent.bind("<KeyRelease>", lambda e: self.schedule_preview_refresh())

        def remove_row(f=row_frame, e=ent):
            self.undangan_inputs["skenario"].remove(e)
            f.destroy()
            self.schedule_preview_refresh()

        btn_del = ctk.CTkButton(row_frame, text="✕", width=30, height=28,
                                fg_color="#EF4444", hover_color="#DC2626", command=remove_row)
        btn_del.grid(row=0, column=1, sticky="e")

        self.undangan_inputs["skenario"].append(ent)

    def update_hari_rapat(self):
        if HAS_TKCALENDAR and "tanggal_rapat" in self.undangan_inputs:
            tanggal = self.undangan_inputs["tanggal_rapat"].get_date()
            hari_nama = ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"]
            hari = hari_nama[tanggal.weekday()]
            self.undangan_inputs["hari_rapat"].configure(state="normal")
            self.undangan_inputs["hari_rapat"].delete(0, tk.END)
            self.undangan_inputs["hari_rapat"].insert(0, hari)
            self.undangan_inputs["hari_rapat"].configure(state="readonly")
            self.schedule_preview_refresh()

    def generate_undangan_paripurna(self):
        if not hasattr(self, 'undangan_inputs'):
            return

        out_dir = filedialog.askdirectory(title="Pilih Folder untuk Menyimpan Surat Undangan")
        if not out_dir:
            return

        template_path = "templates/surat_undangan/paripurna/CONTOH_UNDANGAN_PARIPURNA.docx"
        if not os.path.exists(template_path):
            messagebox.showerror("Error", f"Template tidak ditemukan:\n{template_path}")
            return

        try:
            from docxtpl import DocxTemplate
            import copy

            tanggal_surat = self.format_indonesian_date(self.undangan_inputs["tanggal_surat"].get_date()) if HAS_TKCALENDAR else self.undangan_inputs["tanggal_surat"].get()
            tanggal_rapat = self.format_indonesian_date(self.undangan_inputs["tanggal_rapat"].get_date()) if HAS_TKCALENDAR else self.undangan_inputs["tanggal_rapat"].get()
            hari_rapat = self.undangan_inputs["hari_rapat"].get()
            jam_pelaksanaan = self.undangan_inputs["jam_pelaksanaan"].get()
            pakaian = self.undangan_inputs["pakaian"].get()
            ttd_raw = self.undangan_inputs["penandatanganan"].get()
            jab_ttd, nama_ttd = ttd_raw.split(" - ", 1) if " - " in ttd_raw else ("", ttd_raw)

            skenario_list = [ent.get().strip() for ent in self.undangan_inputs["skenario"] if ent.get().strip()]

            recipients = [
                "Ketua DPRD Kota Bitung",
                "Wakil Ketua DPRD Kota Bitung",
                "Sekretaris DPRD Kota Bitung",
                "Ketua Komisi I DPRD Kota Bitung",
                "Ketua Komisi II DPRD Kota Bitung",
                "Ketua Komisi III DPRD Kota Bitung",
                "Kepala Bagian Umum",
                "Kepala Sub Bagian Keuangan"
            ]

            out_path = os.path.join(out_dir, "Surat_Undangan_Paripurna.docx")

            combined_doc = None
            for i, recipient in enumerate(recipients):
                ctx = {
                    "nomor_undangan": self.undangan_inputs["nomor_undangan"].get(),
                    "tanggal_surat": tanggal_surat,
                    "isi_surat": self.undangan_inputs["isi_surat"].get("1.0", tk.END).strip(),
                    "tanggal_rapat": tanggal_rapat,
                    "hari_rapat": hari_rapat,
                    "jam_pelaksanaan": jam_pelaksanaan,
                    "pakaian": pakaian,
                    "jabatan_ttd": jab_ttd,
                    "nama_ttd": nama_ttd,
                    "penerima": recipient,
                    "skenario1": skenario_list[0] if len(skenario_list) > 0 else "",
                    "skenario2": skenario_list[1] if len(skenario_list) > 1 else "",
                    "skenario3": skenario_list[2] if len(skenario_list) > 2 else "",
                    "skenario4": skenario_list[3] if len(skenario_list) > 3 else "",
                    "skenario5": skenario_list[4] if len(skenario_list) > 4 else "",
                    "skenario6": skenario_list[5] if len(skenario_list) > 5 else "",
                    "skenario7": skenario_list[6] if len(skenario_list) > 6 else "",
                }

                temp_path = os.path.join(out_dir, f"temp_page_{i}.docx")
                doc = DocxTemplate(template_path)
                doc.render(ctx)
                doc.save(temp_path)

            self._combine_undangan_pages(out_dir, out_path, len(recipients), template_path, ctx, recipients)

            for i in range(len(recipients)):
                temp_path = os.path.join(out_dir, f"temp_page_{i}.docx")
                if os.path.exists(temp_path):
                    os.remove(temp_path)

            messagebox.showinfo("Berhasil", f"Surat Undangan Paripurna berhasil dibuat:\n{out_path}\n\nJumlah halaman: {len(recipients)}")

        except Exception as e:
            messagebox.showerror("Error", f"Gagal membuat surat undangan:\n{str(e)}")

    def _combine_undangan_pages(self, out_dir, out_path, num_pages, template_path, base_ctx, recipients):
        from docxtpl import DocxTemplate
        from docx import Document
        from docxcompose.composer import Composer

        first_page_path = os.path.join(out_dir, "combined_first.docx")
        doc = Document(template_path)
        composer = Composer(doc)

        for i in range(num_pages):
            ctx = base_ctx.copy()
            ctx["penerima"] = recipients[i]

            temp_doc = DocxTemplate(template_path)
            temp_doc.render(ctx)
            temp_path = os.path.join(out_dir, f"temp_combine_{i}.docx")
            temp_doc.save(temp_path)

            if i == 0:
                composer = Composer(Document(temp_path))
            else:
                sub_doc = Document(temp_path)
                composer.append(sub_doc)

        composer.save(out_path)

        for i in range(num_pages):
            temp_path = os.path.join(out_dir, f"temp_combine_{i}.docx")
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def generate_undangan_biasa(self):
        messagebox.showinfo("Info", "Fitur Undangan Biasa akan segera tersedia.")

    def load_selected_history(self):
        selected_no = self.combo_history.get()
        if selected_no in self.history_data:
            data = self.history_data[selected_no]
            self.inputs["nomor_surat"].delete(0, tk.END)
            self.inputs["nomor_surat"].insert(0, data.get("nomor_surat", ""))
            self.inputs["nomor_surat_asn"].delete(0, tk.END)
            self.inputs["nomor_surat_asn"].insert(0, data.get("nomor_surat_asn", ""))
            self.inputs["nomor_pemberitahuan_dprd"].delete(0, tk.END)
            self.inputs["nomor_pemberitahuan_dprd"].insert(0, data.get("nomor_pemberitahuan_dprd", ""))
            self.inputs["nomor_pemberitahuan_asn"].delete(0, tk.END)
            self.inputs["nomor_pemberitahuan_asn"].insert(0, data.get("nomor_pemberitahuan_asn", ""))
            self.inputs["nomor_spd_dprd"].delete(0, tk.END)
            self.inputs["nomor_spd_dprd"].insert(0, data.get("nomor_spd_dprd", ""))
            self.inputs["nomor_spd_asn"].delete(0, tk.END)
            self.inputs["nomor_spd_asn"].insert(0, data.get("nomor_spd_asn", ""))

            self.txt_materi_st.delete("1.0", tk.END)
            self.txt_materi_st.insert("1.0", data.get("materi_tugas", ""))
            self.txt_materi_pb.delete("1.0", tk.END)
            self.txt_materi_pb.insert("1.0", data.get("materi_tugas_pb", ""))
            self.txt_dasar_dprd.delete("1.0", tk.END)
            self.txt_dasar_dprd.insert("1.0", data.get("dasar_surat_dprd", ""))
            self.txt_dasar_asn.delete("1.0", tk.END)
            self.txt_dasar_asn.insert("1.0", data.get("dasar_surat_asn", ""))

            self.combo_jenis.set(data.get("jenis_perjalanan", ""))

            tujuan_saved = data.get("tujuan_bertugas_list", [])
            if tujuan_saved:
                self.tujuan_terpilih = tujuan_saved
            else:
                old = data.get("tujuan_bertugas", "")
                self.tujuan_terpilih = [old] if old else []
            self.refresh_tujuan_list_ui()

            for var in self.dprd_vars.values(): var.set(False)
            for var in self.asn_vars.values(): var.set(False)
            for n in data.get("dprd_terpilih", []):
                if n in self.dprd_vars: self.dprd_vars[n].set(True)
            for n in data.get("asn_terpilih", []):
                if n in self.asn_vars: self.asn_vars[n].set(True)

            self.schedule_preview_refresh(immediate=True)
            messagebox.showinfo("Riwayat Dimuat", "Formulir telah diisi dengan data surat sebelumnya.")
        else:
            messagebox.showwarning("Gagal", "Riwayat tidak ditemukan.")

if __name__ == "__main__":
    app = SIPSApp()
    app.mainloop()